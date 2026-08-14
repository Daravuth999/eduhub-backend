"""assessment_tools.py — AI Assessment / Quiz Submission Lab routes.

Student flow: upload a photo/PDF of a completed paper worksheet ->
Gemini extracts the marked answers -> deterministic scoring against the
teacher's answer key -> teacher reviews/corrects in Author Studio ->
teacher awards points (individually or in bulk) via the EXISTING wallet
service, notified via the EXISTING push/notification pipeline. No second
wallet, no fabricated scores, no double-award — see the awarding section.

Collections OWNED by this module (tools/check_collection_ownership.py):
  assessments             — teacher-authored/approved question sets.
  assessment_submissions  — one per student attempt.
  assessment_awards       — one per submission's point-award reservation,
                            unique-indexed on submissionId so a double
                            click (or a retried request) can never award
                            twice — mirrors achievement_tools.py's claim
                            pattern.

Media storage is R2-ONLY (2026-08 explicit product direction — corrected
from an earlier GridFS-fallback draft). Cloudflare R2 is the sole
persistent store for every Assessment Lab binary (student submission
photos/PDFs); MongoDB holds metadata only (assessment definitions,
extracted answers, scoring results, review/award state, audit trail —
see `assessment_schema.py`). There is deliberately NO fallback to GridFS
or any local/Render filesystem: if R2 is unavailable or the upload fails,
`_store_media` raises `SubmissionStorageError` and the route surfaces a
plain, retryable 503 — a file is either genuinely in R2, or the request
failed and nothing was silently stored somewhere else. R2 object keys are
content-addressed (`sha256(bytes)`), so retrying an upload of the exact
same bytes reuses the exact same object instead of accumulating
duplicates, and `_upload_media_to_r2` HEAD-checks before PUT to skip a
redundant write when the object already exists. Deletion is explicit
too — `DELETE /admin/assessments/submissions/{id}` removes the R2 object
and the Mongo doc together (blocked for an already-awarded submission,
to preserve the award's audit trail).

Point-awarding follows achievement_tools.py's claim/credit/finalize
pattern exactly: WalletService.credit(...) is the ONLY thing that ever
moves points, gated by a unique-indexed reservation document plus the
wallet's own idempotency_key — never a second, parallel point system.
Notification follows attendance_tools.py's injected fan_out_push/
build_target_query pattern — title/body are always server-rendered here,
never accepted from the client.
"""
from __future__ import annotations

import asyncio
import hashlib
import io
import logging
import os
from datetime import datetime, timezone

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from pymongo.errors import DuplicateKeyError

import assessment_ai_provider as ai
from assessment_schema import (
    VALID_ASSESSMENT_STATUSES,
    VALID_SUBMISSION_STATUSES,
    build_assessment_document,
    build_award_document,
    build_question,
    build_submission_document,
    new_assessment_id,
    new_award_id,
    new_submission_id,
    normalize_extracted_answer_key,
    normalize_extracted_submission_answers,
    total_points,
    validate_assessment_document,
    validate_submission_document,
)
from assessment_scoring import score_submission

log = logging.getLogger("eduhub.assessment")

COLL_ASSESSMENTS = "assessments"
COLL_SUBMISSIONS = "assessment_submissions"
COLL_AWARDS = "assessment_awards"

# Student submissions: photo or scanned PDF of a completed worksheet.
SUBMISSION_CONTENT_TYPES: dict[str, str] = {
    "image/jpeg": "jpg", "image/jpg": "jpg", "image/png": "png",
    "image/webp": "webp", "image/heic": "heic",
    "application/pdf": "pdf",
}
# Teacher answer-key uploads additionally accept a Word document (text is
# extracted server-side; Gemini never receives raw .docx bytes, which its
# multimodal API does not accept).
ANSWER_KEY_CONTENT_TYPES: dict[str, str] = {
    **SUBMISSION_CONTENT_TYPES,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}
HARD_MAX_MEDIA_BYTES = 25 * 1024 * 1024  # generous for a photographed worksheet page


class SubmissionStorageError(Exception):
    """Raised when R2 could not store an Assessment Lab binary. Routes
    catch this and surface a plain, retryable error — there is no other
    storage backend for this feature to fall back to."""

    def __init__(self, message: str = "") -> None:
        super().__init__(message)
        self.message = message or "Storage is temporarily unavailable. Please try again."


def _r2_config() -> dict | None:
    required = ["R2_ACCOUNT_ID", "R2_ACCESS_KEY_ID", "R2_SECRET_ACCESS_KEY", "R2_BUCKET_NAME", "R2_PUBLIC_URL"]
    cfg = {k: os.environ.get(k, "").strip() for k in required}
    return cfg if all(cfg.values()) else None


def _r2_client(cfg: dict):
    import boto3
    from botocore.config import Config as _BotocoreConfig

    return boto3.client(
        "s3", endpoint_url=f"https://{cfg['R2_ACCOUNT_ID']}.r2.cloudflarestorage.com",
        aws_access_key_id=cfg["R2_ACCESS_KEY_ID"],
        aws_secret_access_key=cfg["R2_SECRET_ACCESS_KEY"],
        region_name="auto",
        config=_BotocoreConfig(signature_version="s3v4"),
    )


def _content_addressed_key(raw: bytes, ext: str, prefix: str) -> str:
    """Deterministic R2 object key = sha256(bytes) — the SAME uploaded
    file always maps to the SAME key, so retrying an upload (or a student
    re-selecting the identical photo) can never create a duplicate R2
    object; it just re-resolves to the object that's already there."""
    digest = hashlib.sha256(raw).hexdigest()
    return f"assessment-media/{prefix}/{digest}.{ext}"


async def _upload_media_to_r2(raw: bytes, key: str, content_type: str) -> str | None:
    """R2-only, NEVER falls back to any other storage. Returns None on any
    failure (env vars absent, boto3 missing, network error) — the CALLER
    (`_store_media`) is responsible for turning that into an honest,
    retryable failure, never a silent write elsewhere. HEAD-checks the
    key first and skips the PUT entirely when the (content-addressed)
    object already exists, so a retried/duplicate upload of identical
    bytes never re-transfers or duplicates storage."""
    cfg = _r2_config()
    if cfg is None:
        return None
    try:
        from botocore.exceptions import ClientError

        def _do_upload() -> bool:
            s3 = _r2_client(cfg)
            try:
                s3.head_object(Bucket=cfg["R2_BUCKET_NAME"], Key=key)
                return True  # already stored — nothing to do
            except ClientError as exc:
                code = str((exc.response or {}).get("Error", {}).get("Code") or "")
                if code not in ("404", "NoSuchKey", "NotFound"):
                    raise
            s3.put_object(Bucket=cfg["R2_BUCKET_NAME"], Key=key, Body=raw, ContentType=content_type)
            return False

        loop = asyncio.get_event_loop()
        already_existed = await loop.run_in_executor(None, _do_upload)
        url = f"{cfg['R2_PUBLIC_URL'].rstrip('/')}/{key}"
        if already_existed:
            log.info("assessment_tools: content-addressed object already exists, skipped upload key=%s", key)
        else:
            log.info("assessment_tools: uploaded %s (%d bytes) url=%s", key, len(raw), url)
        return url
    except Exception:  # noqa: BLE001
        log.exception("assessment_tools: R2 upload failed for key=%s", key)
        return None


async def _delete_media_from_r2(key: str) -> bool:
    """Explicit delete for retention/deletion flows — best-effort: a real
    failure here (network/credentials) is logged CRITICAL for manual
    follow-up rather than blocking the caller, since the Mongo doc removal
    is the user-facing action and a lingering orphaned object is a
    cleanup concern, not a correctness one. Returns True on confirmed
    deletion (including "already gone")."""
    cfg = _r2_config()
    if cfg is None:
        return False
    try:
        def _do_delete():
            _r2_client(cfg).delete_object(Bucket=cfg["R2_BUCKET_NAME"], Key=key)

        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, _do_delete)
        log.info("assessment_tools: deleted R2 object key=%s", key)
        return True
    except Exception:  # noqa: BLE001
        log.critical("assessment_tools: R2 delete FAILED for key=%s — manual cleanup needed", key)
        return False


async def _store_media(raw: bytes, ext: str, content_type: str, prefix: str) -> tuple[str, str]:
    """R2-only. Raises SubmissionStorageError (never falls back to GridFS
    or local disk) if R2 is unavailable or the upload fails. Returns
    (media_ref_url, media_key) — the key is persisted alongside the doc
    so deletion never has to parse it back out of the public URL."""
    key = _content_addressed_key(raw, ext, prefix)
    media_ref = await _upload_media_to_r2(raw, key, content_type)
    if not media_ref:
        raise SubmissionStorageError()
    return media_ref, key


def _extract_docx_text(raw: bytes) -> str:
    """Plain-text extraction from a teacher-uploaded .docx answer key,
    including table cells (a "Quick Answer Key" table is the realistic
    shape for this feature's first fixture) — python-docx is an existing
    project dependency (tuition_receipt_pdf.py's neighbors), not a new one."""
    from docx import Document

    doc = Document(io.BytesIO(raw))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _iso_now() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


async def _sync_award_to_gas(
    clean_id: str, points: float, *,
    gas_url: str, treasury_id: str, treasury_password: str,
) -> tuple[bool, str]:
    """Best-effort treasury->student legacy-balance sync, via the SAME
    proven `action=sendPoints` GAS bridge already used in production by
    Speaking Lab's /points/grant and the referral reward path (server.py) —
    not a new mechanism.

    WHY THIS EXISTS: WalletService.credit() (the caller, before this runs)
    is this app's real, canonical points ledger — it already succeeded by
    the time this is called, and that alone satisfies "the student's
    points were actually credited". But most of the app's student-facing
    UI (the header/dashboard points pill, via usePoints.ts) still polls
    the LEGACY GAS-backed balance by default (REACT_APP_USE_RENDER_POINTS
    is off unless explicitly flagged) — a Mongo-only credit is real but
    invisible there until that flag flips. This call closes that
    visibility gap for the common case; if it fails, the canonical wallet
    credit is NOT reversed (see _award_one) — the outcome is recorded
    honestly (gasSynced=False) instead of silently claimed as done.

    GAS's legacy points ledger is integer-based (Speaking Lab's own grant
    payload sends whole points); an assessment can award fractional points
    (0.5/question), so the amount synced here is rounded to the nearest
    whole point for THIS legacy-visibility leg only — the exact fractional
    amount remains the persisted, authoritative value in points_wallets.

    Never raises. Returns (ok, error_message)."""
    if not (gas_url and treasury_id and treasury_password and clean_id):
        return False, "gas_sync_not_configured"
    amount = max(1, round(points))
    try:
        import httpx

        async with httpx.AsyncClient(
            timeout=httpx.Timeout(12.0, connect=6.0), follow_redirects=True,
        ) as cli:
            r = await cli.post(gas_url, data={
                "action": "sendPoints",
                "id": treasury_id,
                "password": treasury_password,
                "receiverId": clean_id,
                "amount": str(amount),
                "nonce": os.urandom(12).hex(),
            })
        if r.status_code != 200:
            return False, f"HTTP {r.status_code}"
        try:
            j = r.json()
        except Exception:  # noqa: BLE001
            return False, (r.text or "")[:200]
        if isinstance(j, dict) and j.get("success") is True:
            return True, ""
        return False, str((j or {}).get("message") or (j or {}).get("error") or j)[:200]
    except Exception as exc:  # noqa: BLE001
        return False, str(exc)[:200]


async def ensure_assessment_indexes(db) -> None:
    await db[COLL_ASSESSMENTS].create_index("assessmentId", unique=True)
    await db[COLL_ASSESSMENTS].create_index("status")
    await db[COLL_SUBMISSIONS].create_index("submissionId", unique=True)
    await db[COLL_SUBMISSIONS].create_index([("assessmentId", 1), ("studentId", 1)])
    await db[COLL_SUBMISSIONS].create_index("status")
    await db[COLL_AWARDS].create_index("submissionId", unique=True)
    log.info("assessment_tools: indexes ready")


def register_assessment_routes(api: APIRouter, db, require_admin, require_student, *,
                                wallet=None, fan_out_push=None, build_target_query=None,
                                gas_points_login_url=None, gas_treasury_id=None,
                                gas_treasury_password=None) -> None:
    assessments = db[COLL_ASSESSMENTS]
    submissions = db[COLL_SUBMISSIONS]
    awards = db[COLL_AWARDS]

    async def _get_assessment_or_404(assessment_id: str) -> dict:
        doc = await assessments.find_one({"assessmentId": assessment_id}, {"_id": 0})
        if not doc:
            raise HTTPException(404, "Assessment not found.")
        return doc

    async def _get_submission_or_404(submission_id: str) -> dict:
        doc = await submissions.find_one({"submissionId": submission_id}, {"_id": 0})
        if not doc:
            raise HTTPException(404, "Submission not found.")
        return doc

    # ── Teacher: answer-key extraction (review-before-save, nothing persisted here) ──
    @api.post("/admin/assessments/extract-key")
    async def admin_extract_key(file: UploadFile = File(...), admin=Depends(require_admin)):
        _ = admin
        raw = await file.read()
        if not raw:
            raise HTTPException(400, "Uploaded file is empty.")
        if len(raw) > HARD_MAX_MEDIA_BYTES:
            raise HTTPException(413, f"File exceeds the {HARD_MAX_MEDIA_BYTES}-byte limit.")
        content_type = (file.content_type or "").split(";")[0].strip().lower()
        ext = ANSWER_KEY_CONTENT_TYPES.get(content_type)
        if not ext:
            raise HTTPException(415, f"Unsupported content type: {content_type!r}.")

        if ext == "docx":
            try:
                text = _extract_docx_text(raw)
            except Exception as exc:  # noqa: BLE001
                raise HTTPException(400, f"Could not read .docx file: {exc}")
            result = await ai.extract_answer_key(raw_text=text)
        else:
            result = await ai.extract_answer_key(media_bytes=raw, content_type=content_type)

        if not result.get("ok"):
            raise HTTPException(502, f"Answer-key extraction failed: {result.get('reason')}")
        items = normalize_extracted_answer_key(result.get("items") or [])
        return {"ok": True, "engine": result.get("engine"), "questions": items}

    # ── Teacher: assessment CRUD ──────────────────────────────────────────
    @api.post("/admin/assessments")
    async def admin_create_assessment(payload: dict, admin=Depends(require_admin)):
        if not isinstance(payload, dict):
            raise HTTPException(400, "Body must be JSON.")
        title = str(payload.get("title") or "").strip()
        raw_questions = payload.get("questions") or []
        if not title or not isinstance(raw_questions, list) or not raw_questions:
            raise HTTPException(400, "title and a non-empty questions list are required.")
        questions = [
            build_question(
                str(q.get("qid") or f"q{i + 1}"), q.get("prompt"), q.get("correctAnswer"),
                points=q.get("points") or 1.0, choices=q.get("choices"),
            )
            for i, q in enumerate(raw_questions) if isinstance(q, dict)
        ]
        doc = build_assessment_document(
            new_assessment_id(), title, questions,
            subject=payload.get("subject") or "",
            created_by=getattr(admin, "email", "admin"),
            source_ref=payload.get("sourceRef"),
            status="published" if payload.get("publish") else "draft",
            generated_at=_iso_now(),
        )
        ok, errors = validate_assessment_document(doc)
        if not ok:
            raise HTTPException(400, "; ".join(errors))
        await assessments.insert_one(dict(doc))
        doc.pop("_id", None)
        return {"ok": True, "assessment": doc}

    @api.get("/admin/assessments")
    async def admin_list_assessments(admin=Depends(require_admin)):
        _ = admin
        docs = await assessments.find({}, {"_id": 0}).sort("generatedAt", -1).to_list(200)
        return {"ok": True, "assessments": docs}

    @api.get("/admin/assessments/{assessment_id}")
    async def admin_get_assessment(assessment_id: str, admin=Depends(require_admin)):
        _ = admin
        return {"ok": True, "assessment": await _get_assessment_or_404(assessment_id)}

    @api.patch("/admin/assessments/{assessment_id}")
    async def admin_update_assessment(assessment_id: str, payload: dict, admin=Depends(require_admin)):
        _ = admin
        existing = await _get_assessment_or_404(assessment_id)
        if "title" in payload:
            existing["title"] = str(payload["title"] or "").strip()[:200]
        if "questions" in payload and isinstance(payload["questions"], list):
            existing["questions"] = [
                build_question(
                    str(q.get("qid") or f"q{i + 1}"), q.get("prompt"), q.get("correctAnswer"),
                    points=q.get("points") or 1.0, choices=q.get("choices"),
                )
                for i, q in enumerate(payload["questions"]) if isinstance(q, dict)
            ]
            existing["totalPoints"] = total_points(existing["questions"])
        if payload.get("publish"):
            existing["status"] = "published"
        elif payload.get("archive"):
            existing["status"] = "archived"
        elif "status" in payload and payload["status"] in VALID_ASSESSMENT_STATUSES:
            existing["status"] = payload["status"]
        ok, errors = validate_assessment_document(existing)
        if not ok:
            raise HTTPException(400, "; ".join(errors))
        await assessments.update_one({"assessmentId": assessment_id}, {"$set": existing})
        return {"ok": True, "assessment": existing}

    # ── Teacher: submissions review ───────────────────────────────────────
    @api.post("/admin/assessments/{assessment_id}/extraction-check")
    async def admin_extraction_check(assessment_id: str, file: UploadFile = File(...),
                                      admin=Depends(require_admin)):
        """Staging diagnostic switch: runs ONE real Gemini 2.5 Pro read of an
        uploaded worksheet and compares it, per qid, against the
        deterministic mock baseline (the answer key, which is exactly what
        _mock_submission_answers produces). Pure read-only diagnostic —
        nothing is persisted, no submission is created, no points move."""
        _ = admin
        asmt = await _get_assessment_or_404(assessment_id)
        if not ai.ai_available():
            raise HTTPException(
                503, "Real Gemini extraction is not configured in this environment "
                     "(GEMINI_API_KEY missing or mock mode enabled) — run this check on staging.")
        raw = await file.read()
        if not raw:
            raise HTTPException(400, "Uploaded file is empty.")
        if len(raw) > HARD_MAX_MEDIA_BYTES:
            raise HTTPException(413, f"File exceeds the {HARD_MAX_MEDIA_BYTES}-byte limit.")
        content_type = (file.content_type or "").split(";")[0].strip().lower()
        if content_type not in SUBMISSION_CONTENT_TYPES:
            raise HTTPException(415, f"Unsupported content type: {content_type!r}.")

        extraction = await ai.extract_submission_answers(raw, content_type, asmt["questions"])
        if not extraction.get("ok"):
            raise HTTPException(502, f"Real extraction failed: {extraction.get('reason')}")

        known_ids = [q["qid"] for q in asmt["questions"]]
        real = normalize_extracted_submission_answers(
            extraction.get("answers") or [], known_ids, fill_missing=True)
        # Baseline = the assessment's own answer key — exactly what the
        # deterministic mock derives its answers from.
        base_by = {q["qid"]: str(q.get("correctAnswer") or "") for q in asmt["questions"]}
        real_by = {a["qid"]: a for a in real}

        matches = 0
        mismatches, unreadable = [], []
        for q in asmt["questions"]:
            qid = q["qid"]
            r = real_by[qid]
            baseline_answer = base_by.get(qid, "")
            if r["answerState"] != "answered":
                unreadable.append({"qid": qid, "prompt": q.get("prompt"),
                                    "answerState": r["answerState"],
                                    "confidence": r["confidence"],
                                    "baseline": baseline_answer})
            elif str(r["answer"]).strip().casefold() == baseline_answer.strip().casefold():
                matches += 1
            else:
                mismatches.append({"qid": qid, "prompt": q.get("prompt"),
                                    "baseline": baseline_answer, "real": r["answer"],
                                    "confidence": r["confidence"]})

        preview = score_submission(asmt["questions"], real)
        return {
            "ok": True,
            "engine": extraction.get("engine"),
            "model": extraction.get("model"),
            "verification": extraction.get("verification"),
            "total": len(known_ids),
            "matches": matches,
            "mismatches": mismatches,
            "unreadable": unreadable,
            "scorePreview": {
                "correct": preview["correct"], "total": preview["total"],
                "scorePct": preview["scorePct"], "pointsEarned": preview["pointsEarned"],
                "needsReview": preview["needsReview"],
            },
        }

    @api.get("/admin/assessments/{assessment_id}/submissions")
    async def admin_list_submissions(assessment_id: str, status: str | None = None,
                                      admin=Depends(require_admin)):
        _ = admin
        q: dict = {"assessmentId": assessment_id}
        if status:
            q["status"] = status
        docs = await submissions.find(q, {"_id": 0}).sort("submittedAt", -1).to_list(500)
        return {"ok": True, "submissions": docs}

    @api.get("/admin/assessments/submissions/{submission_id}")
    async def admin_get_submission(submission_id: str, admin=Depends(require_admin)):
        _ = admin
        return {"ok": True, "submission": await _get_submission_or_404(submission_id)}

    @api.delete("/admin/assessments/submissions/{submission_id}")
    async def admin_delete_submission(submission_id: str, admin=Depends(require_admin)):
        _ = admin
        sub = await _get_submission_or_404(submission_id)
        if sub.get("status") == "awarded":
            # An awarded submission is the audit trail for a real point
            # credit — deleting it would erase the reason the wallet moved.
            raise HTTPException(409, "Cannot delete a submission whose points were already awarded.")
        media_deleted = False
        key = sub.get("mediaKey")
        if key:
            media_deleted = await _delete_media_from_r2(key)
            if not media_deleted:
                log.critical(
                    "assessment: submission %s deleted from Mongo but its R2 object "
                    "(key=%s) could NOT be confirmed deleted — manual cleanup needed",
                    submission_id, key,
                )
        await submissions.delete_one({"submissionId": submission_id})
        return {"ok": True, "submissionId": submission_id, "mediaDeleted": media_deleted}

    @api.post("/admin/assessments/submissions/{submission_id}/correct")
    async def admin_correct_submission(submission_id: str, payload: dict, admin=Depends(require_admin)):
        sub = await _get_submission_or_404(submission_id)
        if sub.get("status") == "awarded":
            # The credited amount was the persisted score at award time —
            # editing answers afterwards would desync score vs. real credit.
            raise HTTPException(409, "This submission's points were already awarded — its answers are locked.")
        corrections = payload.get("corrections") if isinstance(payload, dict) else None
        if not isinstance(corrections, list):
            raise HTTPException(400, "corrections must be a list of {qid, answer}.")
        asmt = await _get_assessment_or_404(sub["assessmentId"])
        known_ids = {q["qid"] for q in asmt["questions"]}
        admin_email = getattr(admin, "email", "admin")
        now = _iso_now()

        by_qid = {a["qid"]: dict(a) for a in sub.get("extractedAnswers") or []}
        # Original Gemini extraction is preserved forever — backfilled here
        # for submissions that predate the originalExtractedAnswers field.
        original_answers = sub.get("originalExtractedAnswers") or [dict(a) for a in sub.get("extractedAnswers") or []]
        correction_records = list(sub.get("teacherCorrections") or [])
        applied = []
        for c in corrections:
            if not isinstance(c, dict):
                continue
            qid = str(c.get("qid") or "")
            if qid not in known_ids:
                continue
            previous = by_qid.get(qid) or {}
            answer = str(c.get("answer") or "").strip()[:240]
            by_qid[qid] = {"qid": qid, "answer": answer, "confidence": 1.0,
                           "answerState": "answered" if answer else "blank",
                           "source": "teacher"}
            correction_records.append({
                "qid": qid,
                "previousAnswer": previous.get("answer"),
                "previousState": previous.get("answerState"),
                "answer": answer,
                "correctedBy": admin_email,
                "correctedAt": now,
            })
            applied.append(qid)

        order = {q["qid"]: i for i, q in enumerate(asmt["questions"])}
        new_answers = sorted(by_qid.values(), key=lambda a: order.get(a["qid"], len(order)))
        result = score_submission(asmt["questions"], new_answers)
        await submissions.update_one(
            {"submissionId": submission_id},
            {"$set": {
                "extractedAnswers": new_answers,
                "originalExtractedAnswers": original_answers,
                "score": result,
                "status": "reviewed",
                "reviewedAt": now,
                "reviewedBy": admin_email,
                "teacherCorrections": correction_records,
            }},
        )
        return {"ok": True, "score": result, "correctedQids": applied, "status": "reviewed"}

    def _award_summary(award_doc: dict) -> dict:
        """The subset of an award document a teacher/student actually needs
        to trust "AWARDED" means something real — NOT the whole internal
        doc. ``gasSynced`` is explicitly None (not True/False) when no GAS
        bridge was even configured/attempted, so the UI can distinguish
        "not attempted" from "attempted and failed"."""
        return {
            "pointsCredited": award_doc.get("points"),
            "balanceAfter": award_doc.get("balanceAfter"),
            "creditedAt": award_doc.get("creditedAt"),
            "notifiedAt": award_doc.get("notifiedAt"),
            "gasSynced": award_doc.get("gasSynced"),
            "gasSyncError": award_doc.get("gasSyncError"),
        }

    async def _attempt_gas_sync(award_id: str, submission_id: str, clean_id: str, points: float) -> None:
        """Best-effort — see _sync_award_to_gas's docstring for why a
        failure here never blocks or reverses the (already-real)
        WalletService credit. Always persists an honest outcome."""
        ok, err = await _sync_award_to_gas(
            clean_id, points,
            gas_url=gas_points_login_url, treasury_id=gas_treasury_id,
            treasury_password=gas_treasury_password,
        )
        patch = {"gasSynced": ok, "gasSyncError": "" if ok else err, "gasSyncedAt": _iso_now()}
        await awards.update_one({"awardId": award_id}, {"$set": patch})
        # Merge into the submission's `award` subdocument as a single whole-
        # object $set (not a dotted-path update) — a dotted "award.x" $set
        # relies on MongoDB's own nested-path semantics, which the in-memory
        # test fakes used across this codebase's route tests do not
        # replicate (they apply $set via a flat dict.update()).
        sub = await submissions.find_one({"submissionId": submission_id}, {"_id": 0, "award": 1})
        merged_award = {**((sub or {}).get("award") or {}), **patch}
        await submissions.update_one(
            {"submissionId": submission_id},
            {"$set": {"award": merged_award}},
        )
        if not ok:
            log.warning("assessment: legacy GAS balance sync failed for award %s: %s", award_id, err)

    # ── Teacher: award points (individual + bulk), idempotent ───────────
    async def _award_one(submission_id: str, admin_email: str) -> dict:
        if wallet is None:
            return {"submissionId": submission_id, "ok": False, "reason": "wallet_unavailable"}
        sub = await submissions.find_one({"submissionId": submission_id}, {"_id": 0})
        if not sub:
            return {"submissionId": submission_id, "ok": False, "reason": "submission_not_found"}
        if sub.get("status") == "awarded":
            existing = await awards.find_one({"submissionId": submission_id}, {"_id": 0})
            return {"submissionId": submission_id, "ok": True, "duplicate": True,
                     "points": (existing or {}).get("points"),
                     **_award_summary(existing or {})}
        # needs_review is a "teacher should look" signal, never a dead end —
        # any submission with a persisted deterministic score is awardable.
        if sub.get("status") not in ("needs_review", "scored", "reviewed"):
            return {"submissionId": submission_id, "ok": False, "reason": "not_awardable"}
        if not isinstance(sub.get("score"), dict):
            return {"submissionId": submission_id, "ok": False, "reason": "no_persisted_score"}

        score = sub.get("score") or {}
        points = float(score.get("pointsEarned") or 0.0)
        student_id = sub.get("studentId") or ""
        clean_id = sub.get("cleanId") or ""
        award_id = new_award_id()
        award_doc = build_award_document(
            award_id, submission_id, sub["assessmentId"],
            student_id=student_id, clean_id=clean_id, points=points,
            status="pending", generated_at=_iso_now(),
        )
        try:
            await awards.insert_one(dict(award_doc))
        except DuplicateKeyError:
            existing = await awards.find_one({"submissionId": submission_id}, {"_id": 0})
            return {"submissionId": submission_id, "ok": True, "duplicate": True,
                     "points": (existing or {}).get("points"),
                     **_award_summary(existing or {})}

        idem_key = f"assessment_award:{submission_id}"
        try:
            result = await wallet.credit(
                student_id, points,
                source="assessment_award", source_ref=submission_id,
                idempotency_key=idem_key, clean_id=clean_id or None,
                payload={"assessment_id": sub["assessmentId"], "submission_id": submission_id},
            )
        except BaseException as exc:
            await awards.delete_one({"awardId": award_id, "status": "pending"})
            if not isinstance(exc, Exception):
                raise
            log.error("assessment: wallet credit failed for %s: %s", submission_id, exc)
            return {"submissionId": submission_id, "ok": False, "reason": "credit_failed"}

        # From this point on the REAL points_wallets credit has already
        # happened — "AWARDED" below is truthful regardless of what
        # follows (legacy GAS sync / push notification are downstream
        # visibility concerns, never reversed against the wallet).
        balance_after = float((result or {}).get("balance_after") or 0)
        try:
            await awards.update_one(
                {"awardId": award_id},
                {"$set": {"status": "credited", "creditedAt": _iso_now(), "balanceAfter": balance_after}},
            )
            await submissions.update_one(
                {"submissionId": submission_id},
                {"$set": {"status": "awarded", "award": {
                    "pointsCredited": points, "balanceAfter": balance_after,
                    "creditedAt": _iso_now(), "notifiedAt": None,
                    "gasSynced": None, "gasSyncError": None,
                }}},
            )
        except Exception as exc:  # noqa: BLE001
            log.critical(
                "assessment: award %s credited (student=%s submission=%s) but finalize "
                "write failed — manual reconciliation needed: %s",
                award_id, student_id, submission_id, exc,
            )

        # Legacy points-pill visibility bridge — see _sync_award_to_gas's
        # docstring. Best-effort; a failure here is recorded honestly and
        # never blocks or reverses the award above.
        if points > 0 and (clean_id or student_id):
            await _attempt_gas_sync(award_id, submission_id, clean_id or student_id, points)

        if callable(fan_out_push) and callable(build_target_query) and points > 0:
            try:
                title = "Points awarded"
                body = f"You earned {points:g} points for your assessment. Great work!"
                query = build_target_query("students", [clean_id or student_id], None)
                await fan_out_push(query, title, body, "/portal")
                notified_at = _iso_now()
                await awards.update_one({"awardId": award_id}, {"$set": {"notifiedAt": notified_at}})
                sub_now = await submissions.find_one({"submissionId": submission_id}, {"_id": 0, "award": 1})
                merged_award = {**((sub_now or {}).get("award") or {}), "notifiedAt": notified_at}
                await submissions.update_one(
                    {"submissionId": submission_id},
                    {"$set": {"award": merged_award}},
                )
            except Exception as exc:  # noqa: BLE001
                log.warning("assessment: award notification failed for %s: %s", submission_id, exc)

        final_award = await awards.find_one({"awardId": award_id}, {"_id": 0})
        return {"submissionId": submission_id, "ok": True, "duplicate": False,
                 "points": points, "balanceAfter": balance_after,
                 **_award_summary(final_award or {})}

    @api.post("/admin/assessments/submissions/{submission_id}/award")
    async def admin_award_submission(submission_id: str, admin=Depends(require_admin)):
        result = await _award_one(submission_id, getattr(admin, "email", "admin"))
        if not result.get("ok"):
            reason = result.get("reason") or "Award failed."
            conflict = reason in ("already_awarded", "not_awardable", "no_persisted_score")
            raise HTTPException(409 if conflict else 502, reason)
        return result

    @api.post("/admin/assessments/submissions/{submission_id}/retry-gas-sync")
    async def admin_retry_gas_sync(submission_id: str, admin=Depends(require_admin)):
        """Retries ONLY the legacy-visibility GAS bridge for an already-
        awarded submission — never re-credits the wallet (that would be a
        second, real payment). Used when a teacher sees "Wallet credited"
        but "Legacy balance sync: failed" in Author Studio."""
        _ = admin
        sub = await _get_submission_or_404(submission_id)
        if sub.get("status") != "awarded":
            raise HTTPException(409, "This submission has not been awarded yet.")
        award = await awards.find_one({"submissionId": submission_id}, {"_id": 0})
        if not award:
            raise HTTPException(404, "No award record found for this submission.")
        points = float(award.get("points") or 0)
        clean_id = award.get("cleanId") or award.get("studentId") or ""
        if points <= 0 or not clean_id:
            raise HTTPException(400, "Nothing to sync for this award.")
        await _attempt_gas_sync(award["awardId"], submission_id, clean_id, points)
        refreshed = await awards.find_one({"submissionId": submission_id}, {"_id": 0})
        return {"ok": True, "submissionId": submission_id, **_award_summary(refreshed or {})}

    @api.post("/admin/assessments/submissions/bulk-award")
    async def admin_bulk_award(payload: dict, admin=Depends(require_admin)):
        ids = (payload or {}).get("submissionIds")
        if not isinstance(ids, list) or not ids:
            raise HTTPException(400, "submissionIds must be a non-empty list.")
        admin_email = getattr(admin, "email", "admin")
        results = []
        for sid in ids[:200]:
            try:
                results.append(await _award_one(str(sid), admin_email))
            except Exception as exc:  # noqa: BLE001
                log.exception("assessment: bulk award failed for %s", sid)
                results.append({"submissionId": sid, "ok": False, "reason": str(exc)[:200]})
        return {"ok": True, "results": results,
                "awarded": sum(1 for r in results if r.get("ok")),
                "failed": sum(1 for r in results if not r.get("ok"))}

    # ── Student: list, submit, own history ────────────────────────────────
    @api.get("/student/assessments")
    async def student_list_assessments(student=Depends(require_student)):
        docs = await assessments.find({"status": "published"}, {"_id": 0}).sort("generatedAt", -1).to_list(200)
        clean_id = getattr(student, "clean_id", "")
        student_id = getattr(student, "student_id", "")
        for d in docs:
            existing = await submissions.find_one(
                {"assessmentId": d["assessmentId"],
                 "$or": [{"studentId": student_id}, {"cleanId": clean_id}]},
                {"_id": 0, "submissionId": 1, "status": 1, "score": 1},
            )
            d["mySubmission"] = existing
        return {"ok": True, "assessments": docs}

    @api.post("/student/assessments/submit")
    async def student_submit_assessment(assessment_id: str = Form(...), file: UploadFile = File(...),
                                         student=Depends(require_student)):
        asmt = await _get_assessment_or_404(assessment_id)
        if asmt.get("status") != "published":
            raise HTTPException(423, "This assessment is not open for submissions.")

        raw = await file.read()
        if not raw:
            raise HTTPException(400, "Uploaded file is empty.")
        if len(raw) > HARD_MAX_MEDIA_BYTES:
            raise HTTPException(413, f"File exceeds the {HARD_MAX_MEDIA_BYTES}-byte limit.")
        content_type = (file.content_type or "").split(";")[0].strip().lower()
        ext = SUBMISSION_CONTENT_TYPES.get(content_type)
        if not ext:
            raise HTTPException(415, f"Unsupported content type: {content_type!r}.")

        student_id = str(getattr(student, "student_id", "") or "")
        clean_id = str(getattr(student, "clean_id", "") or "")

        # Duplicate-submission guard: the award endpoint is keyed by
        # submissionId, so two independent submission docs for the SAME
        # (assessment, student) attempt could each be awarded separately —
        # a real double-payment risk, not just a UX nuisance. The student
        # PWA already hides the submit action once a non-"failed"
        # submission exists (AssessmentsListPage.jsx's canSubmit check);
        # this enforces the same rule server-side so a retried request or
        # a second client can never create a second live submission.
        existing_submission = await submissions.find_one(
            {"assessmentId": assessment_id,
             "$or": [{"studentId": student_id}, {"cleanId": clean_id}]},
            {"_id": 0, "submissionId": 1, "status": 1},
        )
        if existing_submission and existing_submission.get("status") != "failed":
            raise HTTPException(
                409, "You already submitted this assessment "
                     f"(status: {existing_submission.get('status')}).",
            )

        try:
            media_ref, media_key = await _store_media(raw, ext, content_type, clean_id or student_id or "unknown")
        except SubmissionStorageError as exc:
            # R2-only, on purpose: no GridFS/local-disk fallback. A storage
            # failure must surface as an honest, retryable error — never a
            # silent write to a different backend, and never a submission
            # record referencing a file that doesn't actually exist in R2.
            raise HTTPException(503, exc.message)

        submission_id = new_submission_id()
        doc = build_submission_document(
            submission_id, assessment_id,
            student_id=student_id, clean_id=clean_id,
            media_ref=media_ref, media_key=media_key, content_type=content_type,
            status="processing", generated_at=_iso_now(),
        )
        ok, errors = validate_submission_document(doc)
        if not ok:
            raise HTTPException(500, "; ".join(errors))
        await submissions.insert_one(dict(doc))

        extraction = await ai.extract_submission_answers(raw, content_type, asmt["questions"])
        known_ids = [q["qid"] for q in asmt["questions"]]
        if not extraction.get("ok"):
            await submissions.update_one({"submissionId": submission_id}, {"$set": {"status": "failed"}})
            doc.pop("_id", None)
            doc["status"] = "failed"
            return {"ok": True, "submission": doc, "extractionError": extraction.get("reason")}

        raw_answers = extraction.get("answers") or []
        answers = normalize_extracted_submission_answers(raw_answers, known_ids, fill_missing=True)
        result = score_submission(asmt["questions"], answers)
        status = "needs_review" if result["needsReview"] else "scored"
        # Audit metadata only (per this feature's own data-classification
        # rule: MongoDB holds metadata, never the binary) — an at-a-glance
        # signal for exactly the failure class this exists to catch: did
        # Gemini return nothing (rawAnswerCount=0), did most of it get
        # dropped by the qid whitelist (normalizedAnswerCount << raw), or
        # did everything survive but still score 0 (a real content/
        # vocabulary mismatch, visible per-question in score.details
        # below). Never a second source of truth for the score itself.
        extraction_meta = {
            "engine": extraction.get("engine"),
            "model": extraction.get("model"),
            "extractedAt": _iso_now(),
            "rawAnswerCount": len(raw_answers),
            "normalizedAnswerCount": len(answers),
            "verification": extraction.get("verification"),
        }
        # `originalExtractedAnswers` is the frozen, auditable Gemini result —
        # teacher corrections later rewrite `extractedAnswers`, never this.
        original_answers = [dict(a) for a in answers]
        await submissions.update_one(
            {"submissionId": submission_id},
            {"$set": {"extractedAnswers": answers,
                      "originalExtractedAnswers": original_answers,
                      "score": result, "status": status,
                      "extraction": extraction_meta}},
        )
        doc.pop("_id", None)
        doc.update({"extractedAnswers": answers,
                    "originalExtractedAnswers": original_answers,
                    "score": result, "status": status,
                    "extraction": extraction_meta})
        return {"ok": True, "submission": doc, "engine": extraction.get("engine")}

    @api.get("/student/assessments/submissions")
    async def student_list_submissions(student=Depends(require_student)):
        student_id = str(getattr(student, "student_id", "") or "")
        clean_id = str(getattr(student, "clean_id", "") or "")
        docs = await submissions.find(
            {"$or": [{"studentId": student_id}, {"cleanId": clean_id}]}, {"_id": 0},
        ).sort("submittedAt", -1).to_list(200)
        return {"ok": True, "submissions": docs}

    log.info("assessment_tools: routes registered")
