"""content_parser.py — Smart raw-input → structured-blocks converter.

Takes any raw text (markdown paste, plain prose, or content extracted from
.docx / .txt / .md) and produces the canonical
`chapters: [{title, blocks: [{type, text, ...}]}]` shape the reader expects.

Auto-detected block types:
    heading, paragraph, quote, image, audio, video, embed,
    dialog, mcq, fillblank, transcript, example, markdown(table)

The parser is deliberately permissive: anything it can't confidently classify
stays as a `paragraph`, which renders fine in the reader. No exceptions, no
lost content.
"""
from __future__ import annotations
import re
from typing import Any
from urllib.parse import urlparse


# ---------------- regex table --------------------------------------------- #
RE_H2 = re.compile(r"^##\s+(.+?)\s*$")
RE_H3 = re.compile(r"^###\s+(.+?)\s*$")
RE_QUOTE = re.compile(r"^>\s+(.*)$")
RE_IMG = re.compile(r"^!\[[^\]]*\]\(([^)]+)\)\s*$")
RE_MEDIA_DECLARE = re.compile(r"^(audio|video|embed)\s*:\s*(\S+.*)$", re.I)
RE_DIALOG = re.compile(r"^([A-Za-z][\w\s.'\-]{0,40}?):\s+(.+)$")
RE_OPTION = re.compile(r"^(\*?)\s*([A-Z])[\.\)]\s+(.+)$")
RE_TABLE_ROW = re.compile(r"^\s*\|.*\|\s*$")
RE_HR = re.compile(r"^\s*---+\s*$")
RE_MCQ_Q = re.compile(r"^Q\s*[:\.\-]\s*(.+)$", re.I)
RE_FILL_A = re.compile(r"^A\s*[:\.]\s*(.+)$", re.I)
RE_TRANSCRIPT = re.compile(r"^\[(\d+(?:\.\d+)?)\s*-\s*(\d+(?:\.\d+)?)\]\s*(.+)$")


# ---------------- block helpers ------------------------------------------- #
def _push_paragraph(chapter: dict, buf: list[str]) -> None:
    text = "\n".join(buf).strip()
    if not text:
        return
    chapter["blocks"].append({"type": "paragraph", "text": text})


def _infer_type_from_url(url: str) -> str:
    """Return 'image' | 'audio' | 'video' | 'embed' based on url."""
    low = url.lower().split("?")[0]
    if re.search(r"\.(mp3|m4a|wav|ogg|aac|flac)$", low):
        return "audio"
    if re.search(r"\.(mp4|webm|mov|m3u8)$", low):
        return "video"
    if re.search(r"\.(png|jpe?g|webp|gif|avif|svg)$", low):
        return "image"
    try:
        host = urlparse(url).hostname or ""
        host = host.lower()
        if host.endswith(("youtube.com", "youtu.be", "vimeo.com", "loom.com",
                          "dailymotion.com", "facebook.com")):
            return "embed"
    except Exception:  # noqa: BLE001
        pass
    return "audio"  # safe default for media URLs


# ---------------- core parse ---------------------------------------------- #
def parse_content(raw: str, default_chapter: str = "Main") -> dict[str, Any]:
    """Return {"format": "blocks", "chapters": [...]}.

    The function is pure (no IO). Handles markdown + plain prose + dialog
    + MCQ + fill-blank + media declarations + tables + transcripts.
    """
    if not raw or not isinstance(raw, str):
        return {"format": "blocks", "chapters": []}

    # Normalise CR/LF
    text = raw.replace("\r\n", "\n").replace("\r", "\n").strip("\n")
    lines = text.split("\n")

    chapters: list[dict[str, Any]] = []
    chapter = {"title": default_chapter, "blocks": []}
    para_buf: list[str] = []
    table_buf: list[str] = []

    # State-machine for multi-line constructs
    mcq_state: dict | None = None     # building MCQ: {"question": str, "options": [{letter, text, correct}]}
    dialog_run: bool = False          # currently inside a dialog sequence

    def flush_para() -> None:
        _push_paragraph(chapter, para_buf)
        para_buf.clear()

    def flush_table() -> None:
        if not table_buf:
            return
        chapter["blocks"].append({"type": "markdown", "text": "\n".join(table_buf)})
        table_buf.clear()

    def flush_mcq() -> None:
        nonlocal mcq_state
        if not mcq_state:
            return
        # Collect correct answers
        correct = [o["letter"] for o in mcq_state["options"] if o["correct"]]
        options_text = "\n".join(
            f"{o['letter']}) {o['text']}" for o in mcq_state["options"]
        )
        chapter["blocks"].append({
            "type": "mcq",
            "text": mcq_state["question"],
            "options": options_text,
            "answer": ",".join(correct) if correct else (
                mcq_state["options"][0]["letter"] if mcq_state["options"] else ""
            ),
            "explain": "",
        })
        mcq_state = None

    def new_chapter(title: str) -> None:
        nonlocal chapter, dialog_run
        flush_para()
        flush_table()
        flush_mcq()
        if chapter["blocks"] or chapter["title"] != default_chapter:
            chapters.append(chapter)
        chapter = {"title": title.strip() or default_chapter, "blocks": []}
        dialog_run = False

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- horizontal rule → chapter break -----------------------------
        if RE_HR.match(line):
            new_chapter("Section " + str(len(chapters) + 2))
            i += 1
            continue

        # --- H2 → new chapter --------------------------------------------
        m = RE_H2.match(line)
        if m:
            new_chapter(m.group(1))
            i += 1
            continue

        # --- H3 → heading block ------------------------------------------
        m = RE_H3.match(line)
        if m:
            flush_para()
            flush_table()
            flush_mcq()
            chapter["blocks"].append({"type": "heading", "text": m.group(1).strip()})
            dialog_run = False
            i += 1
            continue

        # --- blank line ---------------------------------------------------
        if stripped == "":
            flush_para()
            flush_table()
            flush_mcq()
            dialog_run = False
            i += 1
            continue

        # --- table --------------------------------------------------------
        if RE_TABLE_ROW.match(line):
            flush_para()
            flush_mcq()
            table_buf.append(line)
            i += 1
            continue
        if table_buf:
            flush_table()

        # --- quote --------------------------------------------------------
        m = RE_QUOTE.match(line)
        if m:
            flush_para()
            flush_mcq()
            chapter["blocks"].append({"type": "quote", "text": m.group(1).strip()})
            i += 1
            continue

        # --- image ![alt](url) -------------------------------------------
        m = RE_IMG.match(line)
        if m:
            flush_para()
            flush_mcq()
            chapter["blocks"].append({"type": "image", "text": m.group(1).strip()})
            i += 1
            continue

        # --- media: `audio: url` / `video: url` / `embed: url` -----------
        m = RE_MEDIA_DECLARE.match(stripped)
        if m:
            flush_para()
            flush_mcq()
            chapter["blocks"].append({
                "type": m.group(1).lower(),
                "text": m.group(2).strip(),
            })
            i += 1
            continue

        # --- transcript "[0.0 - 3.4] Hello there" ------------------------
        m = RE_TRANSCRIPT.match(stripped)
        if m:
            flush_para()
            flush_mcq()
            chapter["blocks"].append({
                "type": "transcript",
                "start": float(m.group(1)),
                "end": float(m.group(2)),
                "text": m.group(3).strip(),
            })
            i += 1
            continue

        # --- MCQ question "Q: ..." ---------------------------------------
        mq = RE_MCQ_Q.match(stripped)
        if mq:
            flush_para()
            flush_mcq()
            mcq_state = {"question": mq.group(1).strip(), "options": []}
            i += 1
            continue

        # --- MCQ option "A) text" / "*A) text" ---------------------------
        if mcq_state is not None:
            mo = RE_OPTION.match(stripped)
            if mo:
                mcq_state["options"].append({
                    "letter": mo.group(2),
                    "text": mo.group(3).strip(),
                    "correct": bool(mo.group(1)),
                })
                i += 1
                continue
            # Fill-in-the-blank right after Q:
            mf = RE_FILL_A.match(stripped)
            if mf:
                chapter["blocks"].append({
                    "type": "fillblank",
                    "text": mcq_state["question"],
                    "answer": mf.group(1).strip(),
                    "explain": "",
                })
                mcq_state = None
                i += 1
                continue
            # Otherwise — close MCQ as a regular paragraph question
            flush_mcq()
            # fall through to other checks

        # --- dialog "Speaker: text" --------------------------------------
        md = RE_DIALOG.match(stripped)
        # Only treat as dialog when it looks like short speaker prefix
        # (<= 3 words, capitalised). This avoids catching regular prose.
        if md and _looks_like_dialog(md.group(1)):
            flush_para()
            flush_table()
            chapter["blocks"].append({
                "type": "dialog",
                "speaker": md.group(1).strip(),
                "text": md.group(2).strip(),
            })
            dialog_run = True
            i += 1
            continue

        # --- default: accumulate into paragraph --------------------------
        # A line continuation — append with a newline so the markdown renderer
        # keeps soft-wrap formatting.
        para_buf.append(stripped)
        dialog_run = False
        i += 1

    # flush tail
    flush_para()
    flush_table()
    flush_mcq()
    if chapter["blocks"] or not chapters:
        chapters.append(chapter)

    # Drop empty chapters
    chapters = [c for c in chapters if c["blocks"]]

    return {"format": "blocks", "chapters": chapters}


def _looks_like_dialog(speaker: str) -> bool:
    """Return True if `speaker` prefix likely identifies a dialog line."""
    if not speaker:
        return False
    words = speaker.strip().split()
    if len(words) > 3:
        return False
    # First word should start with uppercase or be a known role
    first = words[0]
    if not first:
        return False
    # Common role names
    lowered = first.lower()
    if lowered in {"teacher", "student", "narrator", "interviewer",
                   "host", "guest", "doctor", "nurse", "officer"}:
        return True
    if first[0].isupper() and first.isalpha():
        return True
    # Allow "A", "B", "Speaker 1"
    if len(first) == 1 and first.isalpha():
        return True
    return False


# ---------------- docx/txt extract ---------------------------------------- #
def extract_docx(data: bytes) -> str:
    """Convert uploaded .docx bytes to markdown-ish plain text.

    Headings → `##` / `###`, lists → `- `, paragraphs preserved. Uses
    python-docx; tables become markdown tables. No server-side styling
    beyond what the reader supports.
    """
    import io
    from docx import Document  # type: ignore

    doc = Document(io.BytesIO(data))
    out: list[str] = []
    for el in _iter_docx_elements(doc):
        kind, payload = el
        if kind == "para":
            para = payload
            style = (para.style.name or "").lower() if para.style else ""
            text = para.text.rstrip()
            if not text.strip():
                out.append("")
                continue
            if "heading 1" in style or "title" in style:
                out.append(f"## {text}")
            elif "heading 2" in style or "heading" in style:
                out.append(f"### {text}")
            elif "quote" in style:
                out.append(f"> {text}")
            elif "list" in style:
                out.append(f"- {text}")
            else:
                out.append(text)
        elif kind == "table":
            out.extend(_docx_table_to_md(payload))
            out.append("")
    return "\n".join(out)


def _iter_docx_elements(doc):
    """Yield ("para"|"table", element) in document order."""
    from docx.oxml.ns import qn  # type: ignore
    body = doc.element.body
    para_map = {p._element: p for p in doc.paragraphs}
    tbl_map = {t._element: t for t in doc.tables}
    for child in body.iterchildren():
        if child.tag == qn("w:p") and child in para_map:
            yield "para", para_map[child]
        elif child.tag == qn("w:tbl") and child in tbl_map:
            yield "table", tbl_map[child]


def _docx_table_to_md(table) -> list[str]:
    rows: list[str] = []
    for i, row in enumerate(table.rows):
        cells = [c.text.replace("\n", " ").strip() for c in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
    return rows
