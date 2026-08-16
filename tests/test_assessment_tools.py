"""tests/test_assessment_tools.py
====================================
AI Assessment / Quiz Submission Lab — schema, scoring, DOCX extraction, and
route-level tests. Self-contained in-memory fakes (same approach as
tests/test_attendance_checkin.py / tests/test_achievement_claims.py).

The REAL_ANSWER_KEY fixture below is the verbatim content of the first
production test fixture supplied for this feature — a 30-item "Long & Short
Sound Listening Challenge" (word -> LONG/SHORT, 0.5 points each, 15 points
total) — not a synthetic placeholder. Using it end-to-end (normalize ->
validate -> score -> extract-from-docx) proves the pipeline against real
content, matching this feature's "no fake data" requirement.
"""
from __future__ import annotations

import asyncio
import copy
import io
import re

import pytest
from pymongo.errors import DuplicateKeyError

import assessment_ai_provider as ai
import assessment_tools as at
from assessment_scoring import score_submission
from assessment_schema import (
    normalize_extracted_answer_key,
    normalize_extracted_submission_answers,
    total_points,
    validate_assessment_document,
    build_assessment_document,
)


def run(c):
    return asyncio.run(c)


# ── real fixture (verbatim from the supplied teacher answer key) ─────────
REAL_ANSWER_KEY_ITEMS = [
    {"no": 1, "prompt": "sheep", "answer": "LONG", "points": 0.5},
    {"no": 2, "prompt": "ship", "answer": "SHORT", "points": 0.5},
    {"no": 3, "prompt": "cheap", "answer": "LONG", "points": 0.5},
    {"no": 4, "prompt": "chip", "answer": "SHORT", "points": 0.5},
    {"no": 5, "prompt": "seat", "answer": "LONG", "points": 0.5},
    {"no": 6, "prompt": "sit", "answer": "SHORT", "points": 0.5},
    {"no": 7, "prompt": "leave", "answer": "LONG", "points": 0.5},
    {"no": 8, "prompt": "live", "answer": "SHORT", "points": 0.5},
    {"no": 9, "prompt": "reach", "answer": "LONG", "points": 0.5},
    {"no": 10, "prompt": "rich", "answer": "SHORT", "points": 0.5},
    {"no": 11, "prompt": "feel", "answer": "LONG", "points": 0.5},
    {"no": 12, "prompt": "fill", "answer": "SHORT", "points": 0.5},
    {"no": 13, "prompt": "pool", "answer": "LONG", "points": 0.5},
    {"no": 14, "prompt": "pull", "answer": "SHORT", "points": 0.5},
    {"no": 15, "prompt": "fool", "answer": "LONG", "points": 0.5},
    {"no": 16, "prompt": "full", "answer": "SHORT", "points": 0.5},
    {"no": 17, "prompt": "food", "answer": "LONG", "points": 0.5},
    {"no": 18, "prompt": "good", "answer": "SHORT", "points": 0.5},
    {"no": 19, "prompt": "boot", "answer": "LONG", "points": 0.5},
    {"no": 20, "prompt": "book", "answer": "SHORT", "points": 0.5},
    {"no": 21, "prompt": "goose", "answer": "LONG", "points": 0.5},
    {"no": 22, "prompt": "look", "answer": "SHORT", "points": 0.5},
    {"no": 23, "prompt": "feet", "answer": "LONG", "points": 0.5},
    {"no": 24, "prompt": "fit", "answer": "SHORT", "points": 0.5},
    {"no": 25, "prompt": "green", "answer": "LONG", "points": 0.5},
    {"no": 26, "prompt": "grin", "answer": "SHORT", "points": 0.5},
    {"no": 27, "prompt": "peach", "answer": "LONG", "points": 0.5},
    {"no": 28, "prompt": "pitch", "answer": "SHORT", "points": 0.5},
    {"no": 29, "prompt": "sleep", "answer": "LONG", "points": 0.5},
    {"no": 30, "prompt": "slip", "answer": "SHORT", "points": 0.5},
]


# ── schema / scoring / docx tests (no fakes needed) ───────────────────────
def test_normalize_real_answer_key_produces_30_questions_worth_15_points():
    questions = normalize_extracted_answer_key(REAL_ANSWER_KEY_ITEMS)
    assert len(questions) == 30
    assert questions[0] == {"qid": "q1", "prompt": "sheep", "correctAnswer": "LONG", "points": 0.5}
    assert questions[-1]["correctAnswer"] == "SHORT"
    assert total_points(questions) == 15.0
    # alternating LONG/SHORT pattern preserved exactly
    for i, q in enumerate(questions):
        expected = "LONG" if i % 2 == 0 else "SHORT"
        assert q["correctAnswer"] == expected


def test_validate_assessment_document_accepts_real_fixture():
    questions = normalize_extracted_answer_key(REAL_ANSWER_KEY_ITEMS)
    doc = build_assessment_document(
        "asmt_test1", "Long & Short Sound Listening Challenge", questions,
        subject="Phonics", status="published", generated_at="2026-08-12T00:00:00Z",
    )
    ok, errors = validate_assessment_document(doc)
    assert ok, errors
    assert doc["totalPoints"] == 15.0


def test_score_submission_all_correct_yields_full_15_points():
    questions = normalize_extracted_answer_key(REAL_ANSWER_KEY_ITEMS)
    known_ids = [q["qid"] for q in questions]
    # Student answered every question correctly (case/whitespace variance,
    # since a real Gemini extraction of handwriting is never byte-exact).
    raw = [{"qid": q["qid"], "answer": f" {q['correctAnswer'].lower()} ", "confidence": 0.95} for q in questions]
    answers = normalize_extracted_submission_answers(raw, known_ids)
    result = score_submission(questions, answers)
    assert result["correct"] == 30
    assert result["total"] == 30
    assert result["scorePct"] == 100.0
    assert result["pointsEarned"] == 15.0
    assert result["needsReview"] is False


def test_score_submission_partial_and_blank_flags_needs_review():
    questions = normalize_extracted_answer_key(REAL_ANSWER_KEY_ITEMS)
    known_ids = [q["qid"] for q in questions]
    # First 20 correct, last 10 left blank (never extracted at all).
    raw = [{"qid": q["qid"], "answer": q["correctAnswer"], "confidence": 0.9} for q in questions[:20]]
    answers = normalize_extracted_submission_answers(raw, known_ids)
    result = score_submission(questions, answers)
    assert result["correct"] == 20
    assert result["pointsEarned"] == 10.0
    assert result["needsReview"] is True  # blanks trigger review
    blank_detail = result["details"][-1]
    assert blank_detail["givenAnswer"] is None
    assert blank_detail["correct"] is False


def test_normalize_submission_answers_drops_unknown_qid_gemini_invents():
    questions = normalize_extracted_answer_key(REAL_ANSWER_KEY_ITEMS[:3])
    known_ids = [q["qid"] for q in questions]
    raw = [
        {"qid": "q1", "answer": "LONG", "confidence": 0.9},
        {"qid": "q99_invented", "answer": "SHORT", "confidence": 0.9},  # not in known_ids
    ]
    answers = normalize_extracted_submission_answers(raw, known_ids)
    assert len(answers) == 1
    assert answers[0]["qid"] == "q1"


def test_extract_docx_text_pulls_table_rows():
    from docx import Document
    doc = Document()
    doc.add_paragraph("Teacher Script & IPA Answer Key")
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "No."
    table.rows[0].cells[1].text = "Teacher says"
    table.rows[0].cells[2].text = "Answer"
    for item in REAL_ANSWER_KEY_ITEMS[:3]:
        row = table.add_row()
        row.cells[0].text = str(item["no"])
        row.cells[1].text = item["prompt"]
        row.cells[2].text = item["answer"]
    buf = io.BytesIO()
    doc.save(buf)
    text = at._extract_docx_text(buf.getvalue())
    assert "sheep" in text
    assert "LONG" in text
    assert "ship" in text
    assert "SHORT" in text


def test_mock_extract_answer_key_from_text_is_labeled_mock():
    # No GEMINI_API_KEY configured in the test environment -> mock mode,
    # matching video_ai_provider.py's MockVideoProvider precedent.
    result = run(ai.extract_answer_key(raw_text="1. sheep - LONG\n2. ship - SHORT"))
    assert result["ok"] is True
    assert result["engine"] == "mock"
    assert isinstance(result["items"], list) and result["items"]


# ── route-level fakes ──────────────────────────────────────────────────────
def _match(doc, q):
    for k, v in q.items():
        if k == "$or":
            if not any(_match(doc, sub) for sub in v):
                return False
            continue
        dv = doc.get(k)
        if isinstance(v, dict) and "$in" in v:
            if dv not in v["$in"]:
                return False
        elif dv != v:
            return False
    return True


class _Cursor:
    def __init__(self, d):
        self._d = d

    def sort(self, f, d=1):
        self._d.sort(key=lambda x: x.get(f) or "", reverse=(d == -1))
        return self

    async def to_list(self, n):
        return self._d[:n]


class _Coll:
    def __init__(self):
        self.docs = {}
        self._auto = 0

    async def create_index(self, *a, **k):
        return None

    async def find_one(self, q, p=None):
        for d in self.docs.values():
            if _match(d, q):
                o = copy.deepcopy(d)
                if p and p.get("_id") == 0:
                    o.pop("_id", None)
                return o
        return None

    async def insert_one(self, doc):
        key = (doc.get("submissionId") or doc.get("assessmentId") or doc.get("awardId")
               or doc.get("_id") or f"auto{self._auto}")
        self._auto += 1
        doc.setdefault("_id", key)
        self.docs[key] = copy.deepcopy(doc)
        return type("R", (), {"inserted_id": key})()

    async def update_one(self, q, up, upsert=False):
        for d in self.docs.values():
            if _match(d, q):
                if "$set" in up:
                    d.update(up["$set"])
                return type("R", (), {"matched_count": 1})()
        return type("R", (), {"matched_count": 0})()

    async def find_one_and_update(self, q, up, return_document=None, projection=None):
        """Mimics Motor/pymongo's atomic find_one_and_update — the exact
        primitive that makes a filter-guarded compare-and-swap (e.g. "only
        update if correctionVersion still equals what I read") atomic
        against a real MongoDB. This fake is single-threaded/synchronous
        so it cannot reproduce TRUE wall-clock concurrency, but the
        filter-match-then-mutate-in-one-step semantics are real: a second
        caller whose filter no longer matches (because a first caller's
        $set already changed the field) gets None, exactly as it would
        against a real unique document under Mongo's per-document atomicity."""
        for d in self.docs.values():
            if _match(d, q):
                before = copy.deepcopy(d)
                if "$set" in up:
                    d.update(up["$set"])
                after = copy.deepcopy(d)
                result = after if return_document else before
                if projection and projection.get("_id") == 0:
                    result.pop("_id", None)
                return result
        return None

    async def delete_one(self, q):
        for k, d in list(self.docs.items()):
            if _match(d, q):
                del self.docs[k]
                return type("R", (), {"deleted_count": 1})()
        return type("R", (), {"deleted_count": 0})()

    def find(self, q, p=None):
        out = []
        for d in self.docs.values():
            if _match(d, q):
                o = copy.deepcopy(d)
                if p and p.get("_id") == 0:
                    o.pop("_id", None)
                out.append(o)
        return _Cursor(out)


class _AwardsColl(_Coll):
    """Enforces the real unique submissionId index."""

    async def insert_one(self, doc):
        for d in self.docs.values():
            if d.get("submissionId") == doc.get("submissionId"):
                raise DuplicateKeyError("duplicate key: submissionId_1")
        return await super().insert_one(doc)


class _CorrectionsColl(_Coll):
    """Enforces the real sparse-unique clientToken index — this IS the
    idempotency backbone the route relies on for "double-click Apply /
    retry with the same token can never create a second audit record".
    Keyed by `correctionId` (NOT the generic _Coll fallback chain, which
    would pick `submissionId` — wrong here, since one submission can have
    MANY correction docs and they must not overwrite each other)."""

    async def insert_one(self, doc):
        token = doc.get("clientToken")
        if token:
            for d in self.docs.values():
                if d.get("clientToken") == token:
                    raise DuplicateKeyError("duplicate key: clientToken_1")
        key = doc.get("correctionId") or f"auto{self._auto}"
        self._auto += 1
        doc.setdefault("_id", key)
        self.docs[key] = copy.deepcopy(doc)
        return type("R", (), {"inserted_id": key})()


class _DB:
    def __init__(self):
        self._c = {
            at.COLL_ASSESSMENTS: _Coll(),
            at.COLL_SUBMISSIONS: _Coll(),
            at.COLL_AWARDS: _AwardsColl(),
            at.COLL_CORRECTIONS: _CorrectionsColl(),
        }

    def __getitem__(self, name):
        return self._c.setdefault(name, _Coll())


class _Router:
    def __init__(self):
        self.routes = {}

    def get(self, p):
        def d(fn):
            self.routes[("GET", p)] = fn
            return fn
        return d

    def post(self, p):
        def d(fn):
            self.routes[("POST", p)] = fn
            return fn
        return d

    def patch(self, p):
        def d(fn):
            self.routes[("PATCH", p)] = fn
            return fn
        return d

    def delete(self, p):
        def d(fn):
            self.routes[("DELETE", p)] = fn
            return fn
        return d


class _Student:
    def __init__(self, sid="stu_alice", clean="stu094"):
        self.student_id = sid
        self.clean_id = clean


class _Admin:
    email = "teacher@example.com"


class _Wallet:
    """Tracks a REAL running balance per student (not just `amount` echoed
    back) so multi-operation flows — an award followed by one or more
    corrections — can be asserted precisely. For every PRE-EXISTING test,
    each student has exactly one credit from a zero starting balance, so
    `balance_after == amount credited` still holds exactly as before
    (running_balance = 0 + amount) — this is a strict, backward-compatible
    superset of the old fake's behavior, not a behavior change."""

    def __init__(self):
        self.seen = {}
        self.calls = 0
        self.debit_calls = 0
        self.balance_lookups = 0
        self.fail = False
        self.debit_fail = False
        self.get_balance_fail = False
        self.balances = {}

    async def get_balance(self, student_id):
        self.balance_lookups += 1
        if self.get_balance_fail:
            raise RuntimeError("simulated balance lookup failure")
        return float(self.balances.get(student_id, 0.0))

    async def credit(self, student_id, amount, *, source, source_ref=None,
                      idempotency_key=None, clean_id=None, **kw):
        self.calls += 1
        if self.fail:
            raise RuntimeError("simulated wallet failure")
        if idempotency_key is not None and idempotency_key in self.seen:
            return {"ok": True, "duplicate": True, "balance_after": self.seen[idempotency_key]}
        bal = round(float(self.balances.get(student_id, 0.0)) + float(amount), 3)
        self.balances[student_id] = bal
        if idempotency_key is not None:
            self.seen[idempotency_key] = bal
        return {"ok": True, "duplicate": False, "balance_after": bal}

    async def debit(self, student_id, amount, *, source, source_ref=None,
                     idempotency_key=None, clean_id=None, allow_negative=False, **kw):
        self.debit_calls += 1
        if self.debit_fail:
            raise RuntimeError("simulated wallet debit failure")
        if idempotency_key is not None and idempotency_key in self.seen:
            return {"ok": True, "duplicate": True, "balance_after": self.seen[idempotency_key]}
        bal = round(float(self.balances.get(student_id, 0.0)) - float(amount), 3)
        if bal < 0 and not allow_negative:
            raise RuntimeError("insufficient_funds")
        self.balances[student_id] = bal
        if idempotency_key is not None:
            self.seen[idempotency_key] = bal
        return {"ok": True, "duplicate": False, "balance_after": bal}


class _UploadFile:
    def __init__(self, data: bytes, content_type: str):
        self._data = data
        self.content_type = content_type

    async def read(self):
        return self._data


def _call(router, method, path, **kw):
    return run(router.routes[(method, path)](**kw))


def _build(wallet=None, pushes=None):
    db = _DB()
    router = _Router()
    push_log = pushes if pushes is not None else []

    async def fan_out(query, title, body, url):
        push_log.append({"query": query, "title": title, "body": body, "url": url})
        return (1, 0)

    def build_q(target, ids, group):
        return {"target": target, "studentId": {"$in": list(ids or [])}}

    at.register_assessment_routes(
        router, db, lambda: None, lambda: None,
        wallet=wallet, fan_out_push=fan_out, build_target_query=build_q,
    )
    return db, router, push_log


def _patch_media_storage(monkeypatch):
    """Mocks a SUCCESSFUL R2 store for tests that aren't specifically
    about storage behavior (no real R2 credentials exist in the test
    environment). Matches _store_media's real (media_ref, media_key)
    return shape exactly, so every consumer of it downstream works
    unmodified in tests."""
    async def fake_store_media(raw, ext, content_type, prefix):
        key = at._content_addressed_key(raw, ext, prefix)
        return f"https://fake-r2.example/{key}", key

    monkeypatch.setattr(at, "_store_media", fake_store_media)


def _seed_published_assessment(db, questions=None):
    questions = questions or normalize_extracted_answer_key(REAL_ANSWER_KEY_ITEMS)
    doc = build_assessment_document(
        "asmt_fixed", "Long & Short Sound Listening Challenge", questions,
        subject="Phonics", status="published", generated_at="2026-08-12T00:00:00Z",
    )
    db[at.COLL_ASSESSMENTS].docs[doc["assessmentId"]] = dict(doc)
    return doc


# ── route tests ────────────────────────────────────────────────────────────
def test_admin_create_assessment_and_student_sees_only_published(monkeypatch):
    db, router, _ = _build()
    payload = {
        "title": "Long & Short Sound Listening Challenge",
        "subject": "Phonics",
        "questions": REAL_ANSWER_KEY_ITEMS[:2],
        "publish": False,
    }
    # admin_create_assessment expects build_question-shaped keys; adapt.
    payload["questions"] = [
        {"qid": f"q{i+1}", "prompt": it["prompt"], "correctAnswer": it["answer"], "points": it["points"]}
        for i, it in enumerate(REAL_ANSWER_KEY_ITEMS[:2])
    ]
    created = _call(router, "POST", "/admin/assessments", payload=payload, admin=_Admin())
    assert created["ok"] is True
    assert created["assessment"]["status"] == "draft"

    # draft assessment must not appear to students
    listed = _call(router, "GET", "/student/assessments", student=_Student())
    assert listed["assessments"] == []

    _call(router, "PATCH", "/admin/assessments/{assessment_id}",
          assessment_id=created["assessment"]["assessmentId"],
          payload={"publish": True}, admin=_Admin())
    listed = _call(router, "GET", "/student/assessments", student=_Student())
    assert len(listed["assessments"]) == 1
    assert listed["assessments"][0]["mySubmission"] is None


def test_student_submit_scores_against_real_answer_key(monkeypatch):
    _patch_media_storage(monkeypatch)
    db, router, _ = _build()
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        # Simulate a real, mostly-correct extraction: everything right
        # except the very last item is misread as the wrong choice.
        answers = [{"qid": q["qid"], "answer": q["correctAnswer"], "confidence": 0.93} for q in questions[:-1]]
        answers.append({"qid": questions[-1]["qid"], "answer": "LONG", "confidence": 0.4})
        return {"ok": True, "answers": answers, "engine": "mock"}

    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)

    file = _UploadFile(b"fake-jpeg-bytes", "image/jpeg")
    result = _call(
        router, "POST", "/student/assessments/submit",
        assessment_id=asmt["assessmentId"], file=file, student=_Student(),
    )
    assert result["ok"] is True
    sub = result["submission"]
    assert sub["status"] == "needs_review"  # last answer low-confidence
    assert sub["score"]["correct"] == 29
    assert sub["score"]["pointsEarned"] == 14.5
    assert sub["mediaRef"].startswith("https://fake-r2.example/")
    assert sub["mediaKey"]

    # student's own submission history reflects it
    mine = _call(router, "GET", "/student/assessments/submissions", student=_Student())
    assert len(mine["submissions"]) == 1
    assert mine["submissions"][0]["submissionId"] == sub["submissionId"]


def test_student_submit_rejects_unsupported_content_type(monkeypatch):
    _patch_media_storage(monkeypatch)
    db, router, _ = _build()
    asmt = _seed_published_assessment(db)
    file = _UploadFile(b"not-really-a-docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with pytest.raises(Exception) as exc_info:
        _call(router, "POST", "/student/assessments/submit",
              assessment_id=asmt["assessmentId"], file=file, student=_Student())
    assert "415" in str(exc_info.value) or "Unsupported" in str(exc_info.value)


def test_admin_correct_submission_recomputes_score(monkeypatch):
    _patch_media_storage(monkeypatch)
    db, router, _ = _build()
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        # Everything blank at first — an honest "extraction found nothing legible".
        return {"ok": True, "answers": [], "engine": "mock"}

    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)
    file = _UploadFile(b"bytes", "image/png")
    submitted = _call(router, "POST", "/student/assessments/submit",
                       assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub_id = submitted["submission"]["submissionId"]
    assert submitted["submission"]["score"]["correct"] == 0

    corrections = [{"qid": q["qid"], "answer": q["correctAnswer"]} for q in asmt["questions"][:5]]
    corrected = _call(router, "POST", "/admin/assessments/submissions/{submission_id}/correct",
                       submission_id=sub_id,
                       payload={"corrections": corrections}, admin=_Admin())
    assert corrected["ok"] is True
    assert corrected["score"]["correct"] == 5
    assert len(corrected["correctedQids"]) == 5

    stored = db[at.COLL_SUBMISSIONS].docs[sub_id]
    assert stored["status"] == "reviewed"
    assert stored["reviewedBy"] == "teacher@example.com"


def test_award_is_idempotent_on_double_click(monkeypatch):
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    pushes = []
    db, router, pushes = _build(wallet=wallet, pushes=pushes)
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        return {"ok": True, "answers": [{"qid": q["qid"], "answer": q["correctAnswer"], "confidence": 0.9}
                                          for q in questions], "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)

    file = _UploadFile(b"bytes", "image/png")
    submitted = _call(router, "POST", "/student/assessments/submit",
                       assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub_id = submitted["submission"]["submissionId"]
    assert submitted["submission"]["score"]["pointsEarned"] == 15.0

    first = _call(router, "POST", "/admin/assessments/submissions/{submission_id}/award",
                  submission_id=sub_id, admin=_Admin())
    assert first["ok"] is True
    assert first["duplicate"] is False
    assert first["points"] == 15.0
    assert wallet.calls == 1
    assert len(pushes) == 1

    second = _call(router, "POST", "/admin/assessments/submissions/{submission_id}/award",
                    submission_id=sub_id, admin=_Admin())
    assert second["ok"] is True
    assert second["duplicate"] is True
    # A DOUBLE AWARD must never move points twice.
    assert wallet.calls == 1
    assert len(pushes) == 1

    assert db[at.COLL_SUBMISSIONS].docs[sub_id]["status"] == "awarded"


def test_bulk_award_awards_each_submission_once(monkeypatch):
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, pushes = _build(wallet=wallet)
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        return {"ok": True, "answers": [{"qid": q["qid"], "answer": q["correctAnswer"], "confidence": 0.9}
                                          for q in questions], "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)

    sub_ids = []
    for sid in ("stu_a", "stu_b", "stu_c"):
        file = _UploadFile(b"bytes", "image/png")
        submitted = _call(router, "POST", "/student/assessments/submit",
                           assessment_id=asmt["assessmentId"], file=file, student=_Student(sid, sid))
        sub_ids.append(submitted["submission"]["submissionId"])

    result = _call(router, "POST", "/admin/assessments/submissions/bulk-award",
                    payload={"submissionIds": sub_ids}, admin=_Admin())
    assert result["ok"] is True
    assert result["awarded"] == 3
    assert result["failed"] == 0
    assert wallet.calls == 3
    for sid in sub_ids:
        assert db[at.COLL_SUBMISSIONS].docs[sid]["status"] == "awarded"


def test_award_without_wallet_configured_fails_cleanly(monkeypatch):
    _patch_media_storage(monkeypatch)
    db, router, _ = _build(wallet=None)
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        return {"ok": True, "answers": [], "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)

    file = _UploadFile(b"bytes", "image/png")
    submitted = _call(router, "POST", "/student/assessments/submit",
                       assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub_id = submitted["submission"]["submissionId"]

    with pytest.raises(Exception) as exc_info:
        _call(router, "POST", "/admin/assessments/submissions/{submission_id}/award",
              submission_id=sub_id, admin=_Admin())
    assert "wallet_unavailable" in str(exc_info.value) or "502" in str(exc_info.value)


# ── R2-only storage correction (2026-08): no GridFS/local-disk fallback ────
def test_no_gridfs_fallback_exists_anywhere_in_the_module():
    """The module must not even HAVE a GridFS code path any more — not
    just "unused", genuinely removed."""
    assert not hasattr(at, "get_media_bucket")
    assert not hasattr(at, "MEDIA_GRIDFS_BUCKET")
    assert "GridFSBucket" not in open(at.__file__, encoding="utf-8").read()


def test_content_addressed_key_is_deterministic_by_bytes_and_scoped_by_prefix():
    raw_a = b"identical worksheet photo bytes"
    raw_b = b"a completely different photo"
    key1 = at._content_addressed_key(raw_a, "png", "stu094")
    key2 = at._content_addressed_key(raw_a, "png", "stu094")
    assert key1 == key2  # same bytes, same prefix -> same key every time

    key_other_student = at._content_addressed_key(raw_a, "png", "stu095")
    assert key_other_student != key1  # scoped by prefix, not a global dedupe

    key_diff_bytes = at._content_addressed_key(raw_b, "png", "stu094")
    assert key_diff_bytes != key1  # different content -> different key
    assert key1.startswith("assessment-media/stu094/")
    assert key1.endswith(".png")


def test_store_media_raises_and_never_falls_back_when_r2_unavailable(monkeypatch):
    async def fake_upload_returns_none(raw, key, content_type):
        return None
    monkeypatch.setattr(at, "_upload_media_to_r2", fake_upload_returns_none)

    with pytest.raises(at.SubmissionStorageError):
        run(at._store_media(b"raw-bytes", "png", "image/png", "stu094"))


def test_submit_route_fails_honestly_with_no_fallback_when_r2_unavailable(monkeypatch):
    """R2 down (or unconfigured) must surface as a clean, retryable error —
    and must NEVER create an orphaned submission doc referencing a file
    that doesn't actually exist anywhere."""
    db, router, _ = _build()
    asmt = _seed_published_assessment(db)

    async def fake_upload_returns_none(raw, key, content_type):
        return None
    monkeypatch.setattr(at, "_upload_media_to_r2", fake_upload_returns_none)

    file = _UploadFile(b"worksheet-bytes", "image/png")
    with pytest.raises(Exception) as exc_info:
        _call(router, "POST", "/student/assessments/submit",
              assessment_id=asmt["assessmentId"], file=file, student=_Student())
    assert "503" in str(exc_info.value) or "temporarily unavailable" in str(exc_info.value)
    # No fallback write happened anywhere — the submission collection is
    # exactly as empty as before the failed attempt.
    assert db[at.COLL_SUBMISSIONS].docs == {}


def test_upload_to_r2_skips_put_when_content_addressed_object_already_exists(monkeypatch):
    """Proves the HEAD-before-PUT dedupe: retrying an upload of IDENTICAL
    bytes must not re-transfer/duplicate the object."""
    from botocore.exceptions import ClientError

    monkeypatch.setenv("R2_ACCOUNT_ID", "acct")
    monkeypatch.setenv("R2_ACCESS_KEY_ID", "key")
    monkeypatch.setenv("R2_SECRET_ACCESS_KEY", "secret")
    monkeypatch.setenv("R2_BUCKET_NAME", "bucket")
    monkeypatch.setenv("R2_PUBLIC_URL", "https://cdn.example.com")

    calls = {"head": 0, "put": 0}

    class _FakeClientAlreadyExists:
        def head_object(self, **kw):
            calls["head"] += 1
            return {}  # exists — no ClientError raised

        def put_object(self, **kw):
            calls["put"] += 1

    import boto3
    monkeypatch.setattr(boto3, "client", lambda *a, **kw: _FakeClientAlreadyExists())

    url = run(at._upload_media_to_r2(b"bytes", "assessment-media/stu094/abc.png", "image/png"))
    assert url == "https://cdn.example.com/assessment-media/stu094/abc.png"
    assert calls["head"] == 1
    assert calls["put"] == 0  # skipped — object already there, not duplicated


def test_upload_to_r2_puts_when_object_does_not_yet_exist(monkeypatch):
    from botocore.exceptions import ClientError

    monkeypatch.setenv("R2_ACCOUNT_ID", "acct")
    monkeypatch.setenv("R2_ACCESS_KEY_ID", "key")
    monkeypatch.setenv("R2_SECRET_ACCESS_KEY", "secret")
    monkeypatch.setenv("R2_BUCKET_NAME", "bucket")
    monkeypatch.setenv("R2_PUBLIC_URL", "https://cdn.example.com")

    calls = {"head": 0, "put": 0}

    class _FakeClientNotFound:
        def head_object(self, **kw):
            calls["head"] += 1
            raise ClientError({"Error": {"Code": "404", "Message": "Not Found"}}, "HeadObject")

        def put_object(self, **kw):
            calls["put"] += 1

    import boto3
    monkeypatch.setattr(boto3, "client", lambda *a, **kw: _FakeClientNotFound())

    url = run(at._upload_media_to_r2(b"bytes", "assessment-media/stu094/new.png", "image/png"))
    assert url == "https://cdn.example.com/assessment-media/stu094/new.png"
    assert calls["head"] == 1
    assert calls["put"] == 1


def test_delete_submission_deletes_r2_object_and_mongo_doc(monkeypatch):
    _patch_media_storage(monkeypatch)
    db, router, _ = _build()
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        return {"ok": True, "answers": [], "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)

    file = _UploadFile(b"bytes", "image/png")
    submitted = _call(router, "POST", "/student/assessments/submit",
                       assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub_id = submitted["submission"]["submissionId"]
    assert sub_id in db[at.COLL_SUBMISSIONS].docs

    deleted_keys = []
    async def fake_delete(key):
        deleted_keys.append(key)
        return True
    monkeypatch.setattr(at, "_delete_media_from_r2", fake_delete)

    result = _call(router, "DELETE", "/admin/assessments/submissions/{submission_id}",
                    submission_id=sub_id, admin=_Admin())
    assert result["ok"] is True
    assert result["mediaDeleted"] is True
    assert deleted_keys == [submitted["submission"]["mediaKey"]]
    assert sub_id not in db[at.COLL_SUBMISSIONS].docs


def test_delete_submission_blocked_once_already_awarded(monkeypatch):
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, _ = _build(wallet=wallet)
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        return {"ok": True, "answers": [{"qid": q["qid"], "answer": q["correctAnswer"], "confidence": 0.9}
                                          for q in questions], "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)

    file = _UploadFile(b"bytes", "image/png")
    submitted = _call(router, "POST", "/student/assessments/submit",
                       assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub_id = submitted["submission"]["submissionId"]
    _call(router, "POST", "/admin/assessments/submissions/{submission_id}/award",
          submission_id=sub_id, admin=_Admin())
    assert db[at.COLL_SUBMISSIONS].docs[sub_id]["status"] == "awarded"

    with pytest.raises(Exception) as exc_info:
        _call(router, "DELETE", "/admin/assessments/submissions/{submission_id}",
              submission_id=sub_id, admin=_Admin())
    assert "409" in str(exc_info.value) or "already been awarded" in str(exc_info.value) or "already awarded" in str(exc_info.value)
    # The submission — the audit trail for the real point credit — must survive.
    assert sub_id in db[at.COLL_SUBMISSIONS].docs


# ── 2026-08 production incident regression: real 15/30 -> 7.5/15 flow ─────
# A real student submission on the live "Long & Short Sound Listening
# Challenge" assessment scored 0/30 despite a genuinely ~15/30-correct
# worksheet. This suite proves the FULL lifecycle — submit -> score ->
# Author Studio visibility -> award -> idempotent double-award -> the new
# extraction diagnostic metadata — using a MIXED (not all-correct,
# not all-wrong) extraction result built from the real answer key, exactly
# matching the reported scenario's shape.
def _mixed_extraction_15_of_30(questions):
    """15 correct, 15 wrong — flips every SECOND question's answer to the
    other real vocabulary value (LONG<->SHORT), never a placeholder."""
    out = []
    for i, q in enumerate(questions):
        correct = q["correctAnswer"]
        given = correct if i % 2 == 0 else ("SHORT" if correct == "LONG" else "LONG")
        out.append({"qid": q["qid"], "answer": given, "confidence": 0.93})
    return out


def test_real_15_of_30_submission_flows_through_the_full_lifecycle(monkeypatch):
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, pushes = _build(wallet=wallet)
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        return {"ok": True, "answers": _mixed_extraction_15_of_30(questions), "engine": "gemini"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)

    file = _UploadFile(b"real-worksheet-bytes", "image/jpeg")
    submitted = _call(router, "POST", "/student/assessments/submit",
                       assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub = submitted["submission"]
    assert sub["score"]["correct"] == 15
    assert sub["score"]["total"] == 30
    assert sub["score"]["pointsEarned"] == 7.5
    assert sub["status"] == "scored"
    # New diagnostic: Gemini genuinely returned 30 answers and all 30
    # survived the qid whitelist — distinguishes "extracted but half
    # wrong" from "extracted nothing" for exactly this failure class.
    # (Meta now also persists model/extractedAt/verification — audit trail.)
    assert sub["extraction"]["engine"] == "gemini"
    assert sub["extraction"]["rawAnswerCount"] == 30
    assert sub["extraction"]["normalizedAnswerCount"] == 30
    assert "extractedAt" in sub["extraction"]

    # Visible to Author Studio via the SAME query path production uses.
    listed = _call(router, "GET", "/admin/assessments/{assessment_id}/submissions",
                    assessment_id=asmt["assessmentId"], status=None, admin=_Admin())
    assert len(listed["submissions"]) == 1
    assert listed["submissions"][0]["score"]["pointsEarned"] == 7.5
    # Per-question breakdown is real and persisted, not fabricated —
    # exactly what Author Studio needs to render a given-vs-correct table.
    details = listed["submissions"][0]["score"]["details"]
    assert len(details) == 30
    assert sum(1 for d in details if d["correct"]) == 15

    sub_id = sub["submissionId"]
    award = _call(router, "POST", "/admin/assessments/submissions/{submission_id}/award",
                  submission_id=sub_id, admin=_Admin())
    assert award["points"] == 7.5
    assert wallet.calls == 1
    assert len(pushes) == 1

    # Idempotent: a second click must NOT award a second time.
    award2 = _call(router, "POST", "/admin/assessments/submissions/{submission_id}/award",
                   submission_id=sub_id, admin=_Admin())
    assert award2["duplicate"] is True
    assert wallet.calls == 1
    assert len(pushes) == 1


def test_extraction_failure_never_fabricates_a_score(monkeypatch):
    """Case: Gemini extraction itself fails (ok:false). Must land on
    status='failed' with NO score field — never a fabricated 0 that looks
    like a real (but wrong) grading result."""
    _patch_media_storage(monkeypatch)
    db, router, _ = _build()
    asmt = _seed_published_assessment(db)

    async def fake_extract_fails(media_bytes, content_type, questions):
        return {"ok": False, "reason": "provider_rejected: Gemini HTTP 500"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract_fails)

    file = _UploadFile(b"bytes", "image/png")
    result = _call(router, "POST", "/student/assessments/submit",
                    assessment_id=asmt["assessmentId"], file=file, student=_Student())
    assert result["submission"]["status"] == "failed"
    assert result["submission"].get("score") is None
    assert "extractionError" in result
    stored = db[at.COLL_SUBMISSIONS].docs[result["submission"]["submissionId"]]
    assert stored["status"] == "failed"
    # The real reason must survive past this one response — persisted on
    # the document itself, so a later visit (student re-opening the
    # assessment, or a teacher/admin looking at the row in Author Studio)
    # can see WHY it failed instead of a bare, undiagnosable "failed".
    assert stored["extractionError"] == "provider_rejected: Gemini HTTP 500"
    assert result["submission"]["extractionError"] == "provider_rejected: Gemini HTTP 500"

    # Author Studio's own list route must surface it too — not just the
    # one-shot submit response.
    listed = _call(router, "GET", "/admin/assessments/{assessment_id}/submissions",
                    assessment_id=asmt["assessmentId"], status=None, admin=_Admin())
    assert listed["submissions"][0]["extractionError"] == "provider_rejected: Gemini HTTP 500"


def test_submitting_to_a_nonexistent_assessment_is_rejected_honestly(monkeypatch):
    _patch_media_storage(monkeypatch)
    db, router, _ = _build()
    file = _UploadFile(b"bytes", "image/png")
    with pytest.raises(Exception) as exc_info:
        _call(router, "POST", "/student/assessments/submit",
              assessment_id="asmt_does_not_exist", file=file, student=_Student())
    assert "404" in str(exc_info.value) or "not found" in str(exc_info.value).lower()


def test_duplicate_submission_is_rejected_before_creating_a_second_live_record(monkeypatch):
    """A double-tap or retried request must never create a second
    submission for the same (assessment, student) — the award endpoint is
    keyed by submissionId, so two live submissions would each be
    independently awardable, a real double-payment risk."""
    _patch_media_storage(monkeypatch)
    db, router, _ = _build()
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        return {"ok": True, "answers": [{"qid": q["qid"], "answer": q["correctAnswer"], "confidence": 0.9}
                                          for q in questions], "engine": "gemini"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)

    file1 = _UploadFile(b"bytes", "image/png")
    first = _call(router, "POST", "/student/assessments/submit",
                  assessment_id=asmt["assessmentId"], file=file1, student=_Student())
    assert first["submission"]["status"] == "scored"

    file2 = _UploadFile(b"bytes2", "image/png")
    with pytest.raises(Exception) as exc_info:
        _call(router, "POST", "/student/assessments/submit",
              assessment_id=asmt["assessmentId"], file=file2, student=_Student())
    assert "409" in str(exc_info.value) or "already submitted" in str(exc_info.value).lower()
    # Exactly one submission exists for this (assessment, student).
    assert len(db[at.COLL_SUBMISSIONS].docs) == 1


def test_resubmission_after_a_failed_extraction_is_allowed(monkeypatch):
    """A 'failed' submission is the one case resubmission must be allowed
    — the student never got a real result the first time."""
    _patch_media_storage(monkeypatch)
    db, router, _ = _build()
    asmt = _seed_published_assessment(db)

    call_count = {"n": 0}

    async def fake_extract(media_bytes, content_type, questions):
        call_count["n"] += 1
        if call_count["n"] == 1:
            return {"ok": False, "reason": "bad_response"}
        return {"ok": True, "answers": [{"qid": q["qid"], "answer": q["correctAnswer"], "confidence": 0.9}
                                          for q in questions], "engine": "gemini"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)

    file1 = _UploadFile(b"bytes", "image/png")
    first = _call(router, "POST", "/student/assessments/submit",
                  assessment_id=asmt["assessmentId"], file=file1, student=_Student())
    assert first["submission"]["status"] == "failed"

    file2 = _UploadFile(b"bytes2", "image/png")
    second = _call(router, "POST", "/student/assessments/submit",
                   assessment_id=asmt["assessmentId"], file=file2, student=_Student())
    assert second["submission"]["status"] == "scored"
    assert len(db[at.COLL_SUBMISSIONS].docs) == 2


def test_vocabulary_mismatched_extraction_scores_zero_but_diagnostic_reveals_data_was_received(monkeypatch):
    """The exact shape of the reported production incident: Gemini
    confidently (high confidence, every qid present) extracts something
    that doesn't match the answer key's vocabulary at all — must score 0
    honestly (never inflated), but the extraction diagnostic must prove
    data WAS received (rawAnswerCount=30), distinguishing this from "Gemini
    returned nothing"."""
    _patch_media_storage(monkeypatch)
    db, router, _ = _build()
    asmt = _seed_published_assessment(db)

    async def fake_extract_wrong_vocab(media_bytes, content_type, questions):
        # Echoes the PROMPT word instead of the LONG/SHORT classification —
        # the exact failure mode _answer_vocabulary_hint now guards against.
        return {"ok": True, "answers": [{"qid": q["qid"], "answer": q["prompt"], "confidence": 0.95}
                                          for q in questions], "engine": "gemini"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract_wrong_vocab)

    file = _UploadFile(b"bytes", "image/png")
    result = _call(router, "POST", "/student/assessments/submit",
                    assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub = result["submission"]
    assert sub["score"]["correct"] == 0
    assert sub["status"] == "scored"  # high confidence throughout -> not needs_review
    assert sub["extraction"]["rawAnswerCount"] == 30
    assert sub["extraction"]["normalizedAnswerCount"] == 30
    # The per-question detail proves exactly what was wrong — visible data,
    # not a silent zero.
    first_detail = sub["score"]["details"][0]
    assert first_detail["givenAnswer"] == "sheep"
    assert first_detail["correctAnswer"] == "LONG"


# ── production incident: "AWARDED" shown but student's visible points ─────
# never moved (2026-08). Root cause: WalletService.credit() — the real,
# canonical points ledger — succeeded correctly (confirmed by inspecting
# it against achievement_tools.py's identical, proven pattern), but the
# DEFAULT student-facing points display (usePoints.ts) polls the LEGACY
# GAS-backed balance, not points_wallets, unless REACT_APP_USE_RENDER_POINTS
# is explicitly flagged on. A Mongo-only credit is real but invisible
# there. These tests cover the fix: a best-effort legacy-visibility bridge
# (the SAME proven action=sendPoints treasury->student GAS call already
# used in production by Speaking Lab's /points/grant), which — critically —
# never blocks or reverses the wallet credit if it fails, and is always
# recorded honestly rather than silently assumed to have worked.
def _build_with_gas(monkeypatch, wallet=None, gas_ok=True, gas_error="GAS unreachable"):
    db, router, pushes = _build(wallet=wallet)
    calls = []

    async def fake_gas_sync(clean_id, points, *, gas_url, treasury_id, treasury_password):
        calls.append({"clean_id": clean_id, "points": points})
        return (gas_ok, "" if gas_ok else gas_error)

    monkeypatch.setattr(at, "_sync_award_to_gas", fake_gas_sync)
    return db, router, pushes, calls


def test_award_success_persists_wallet_proof_and_gas_sync_outcome_on_the_submission(monkeypatch):
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, pushes, gas_calls = _build_with_gas(monkeypatch, wallet=wallet, gas_ok=True)
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        return {"ok": True, "answers": [{"qid": q["qid"], "answer": q["correctAnswer"], "confidence": 0.9}
                                          for q in questions], "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)

    file = _UploadFile(b"bytes", "image/png")
    submitted = _call(router, "POST", "/student/assessments/submit",
                       assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub_id = submitted["submission"]["submissionId"]

    result = _call(router, "POST", "/admin/assessments/submissions/{submission_id}/award",
                    submission_id=sub_id, admin=_Admin())
    assert result["ok"] is True
    assert result["pointsCredited"] == 15.0
    assert result["balanceAfter"] == 15.0  # _Wallet's fake balance_after == amount credited
    assert result["gasSynced"] is True
    assert len(gas_calls) == 1
    assert gas_calls[0]["clean_id"] == "stu094"  # _Student()'s clean_id
    assert gas_calls[0]["points"] == 15.0
    assert len(pushes) == 1

    stored = db[at.COLL_SUBMISSIONS].docs[sub_id]
    assert stored["status"] == "awarded"
    assert stored["award"]["pointsCredited"] == 15.0
    assert stored["award"]["balanceAfter"] == 15.0
    assert stored["award"]["gasSynced"] is True
    assert stored["award"]["notifiedAt"] is not None


def test_gas_sync_failure_never_blocks_or_reverses_the_real_wallet_credit(monkeypatch):
    """The exact reported production symptom's honest resolution: the
    wallet credit (the real, canonical points mutation) already succeeded
    by the time the legacy GAS bridge is even attempted — its failure
    must never undo that, never raise, and never silently claim success."""
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, pushes, gas_calls = _build_with_gas(
        monkeypatch, wallet=wallet, gas_ok=False, gas_error="GAS unreachable",
    )
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        return {"ok": True, "answers": [{"qid": q["qid"], "answer": q["correctAnswer"], "confidence": 0.9}
                                          for q in questions], "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)

    file = _UploadFile(b"bytes", "image/png")
    submitted = _call(router, "POST", "/student/assessments/submit",
                       assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub_id = submitted["submission"]["submissionId"]

    result = _call(router, "POST", "/admin/assessments/submissions/{submission_id}/award",
                    submission_id=sub_id, admin=_Admin())
    # The award itself is still successful — the REAL ledger (WalletService)
    # was credited; only the legacy-visibility bridge failed.
    assert result["ok"] is True
    assert wallet.calls == 1
    assert result["balanceAfter"] == 15.0
    assert result["gasSynced"] is False
    assert "GAS unreachable" in result["gasSyncError"]

    stored = db[at.COLL_SUBMISSIONS].docs[sub_id]
    assert stored["status"] == "awarded"  # never blocked by the GAS failure
    assert stored["award"]["gasSynced"] is False
    assert "GAS unreachable" in stored["award"]["gasSyncError"]
    # A push notification is STILL sent — the real credit succeeded.
    assert len(pushes) == 1


def test_retry_gas_sync_resyncs_without_ever_recrediting_the_wallet(monkeypatch):
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, pushes, gas_calls = _build_with_gas(monkeypatch, wallet=wallet, gas_ok=False)
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        return {"ok": True, "answers": [{"qid": q["qid"], "answer": q["correctAnswer"], "confidence": 0.9}
                                          for q in questions], "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)

    file = _UploadFile(b"bytes", "image/png")
    submitted = _call(router, "POST", "/student/assessments/submit",
                       assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub_id = submitted["submission"]["submissionId"]

    first = _call(router, "POST", "/admin/assessments/submissions/{submission_id}/award",
                  submission_id=sub_id, admin=_Admin())
    assert first["gasSynced"] is False
    assert wallet.calls == 1

    # GAS recovers — teacher clicks "Retry sync".
    async def fake_gas_sync_now_ok(clean_id, points, *, gas_url, treasury_id, treasury_password):
        gas_calls.append({"clean_id": clean_id, "points": points})
        return (True, "")
    monkeypatch.setattr(at, "_sync_award_to_gas", fake_gas_sync_now_ok)

    retried = _call(router, "POST", "/admin/assessments/submissions/{submission_id}/retry-gas-sync",
                     submission_id=sub_id, admin=_Admin())
    assert retried["ok"] is True
    assert retried["gasSynced"] is True
    # The wallet must NEVER be credited a second time by a sync retry.
    assert wallet.calls == 1

    stored = db[at.COLL_SUBMISSIONS].docs[sub_id]
    assert stored["award"]["gasSynced"] is True


def test_retry_gas_sync_rejects_a_submission_that_was_never_awarded(monkeypatch):
    _patch_media_storage(monkeypatch)
    db, router, pushes, gas_calls = _build_with_gas(monkeypatch, wallet=_Wallet())
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        return {"ok": True, "answers": [], "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)

    file = _UploadFile(b"bytes", "image/png")
    submitted = _call(router, "POST", "/student/assessments/submit",
                       assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub_id = submitted["submission"]["submissionId"]

    with pytest.raises(Exception) as exc_info:
        _call(router, "POST", "/admin/assessments/submissions/{submission_id}/retry-gas-sync",
              submission_id=sub_id, admin=_Admin())
    assert "409" in str(exc_info.value) or "not been awarded" in str(exc_info.value)


def test_admin_submission_list_resolves_real_student_display_name(monkeypatch):
    """Author Studio's submission list must show WHO a worksheet belongs
    to — a raw studentId/cleanId code is not enough for a teacher to
    recognize a student. Resolved from db.students (the same collection
    attendance_tools.py/tuition_tools.py already join against), batched
    for the whole page — never a fabricated or guessed name."""
    _patch_media_storage(monkeypatch)
    db, router, pushes = _build(wallet=_Wallet())
    asmt = _seed_published_assessment(db)
    db["students"].docs["s1"] = {
        "student_id": "stu_alice", "clean_id": "stu094", "display_name": "Alice Chan",
    }

    async def fake_extract(media_bytes, content_type, questions):
        return {"ok": True, "answers": [], "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)

    file = _UploadFile(b"bytes", "image/png")
    _call(router, "POST", "/student/assessments/submit",
          assessment_id=asmt["assessmentId"], file=file, student=_Student())

    listed = _call(router, "GET", "/admin/assessments/{assessment_id}/submissions",
                    assessment_id=asmt["assessmentId"], status=None, admin=_Admin())
    assert len(listed["submissions"]) == 1
    assert listed["submissions"][0]["studentName"] == "Alice Chan"


# ── post-award correction (reverse review) ────────────────────────────────
def _award_full_submission(db, router, wallet, *, sid="stu_alice", clean="stu094"):
    """Submits a perfect (15/15, 15.0 pt) worksheet for `sid` and awards
    it, returning (submission_id, questions). Shared setup for every
    correction test below — the correction itself is each test's subject,
    not the award path (already covered above)."""
    asmt = db[at.COLL_ASSESSMENTS].docs.get("asmt_fixed") or _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        return {"ok": True, "answers": [{"qid": q["qid"], "answer": q["correctAnswer"], "confidence": 0.95}
                                          for q in questions], "engine": "mock"}
    import unittest.mock as _mock
    with _mock.patch.object(ai, "extract_submission_answers", fake_extract):
        file = _UploadFile(b"bytes", "image/png")
        submitted = _call(router, "POST", "/student/assessments/submit",
                           assessment_id=asmt["assessmentId"], file=file, student=_Student(sid, clean))
    sub_id = submitted["submission"]["submissionId"]
    awarded = _call(router, "POST", "/admin/assessments/submissions/{submission_id}/award",
                     submission_id=sub_id, admin=_Admin())
    assert awarded["ok"] is True and awarded["points"] == 15.0
    return sub_id, asmt["questions"]


def _correct(router, sub_id, *, corrections, reason="student_evidence_accepted",
             reason_note="", client_token="tok1", expected_version=0):
    return _call(router, "POST", "/admin/assessments/submissions/{submission_id}/correction",
                 submission_id=sub_id,
                 payload={"corrections": corrections, "reason": reason, "reasonNote": reason_note,
                          "clientToken": client_token, "expectedVersion": expected_version},
                 admin=_Admin())


def test_correction_route_requires_admin_dependency():
    """Permission enforcement — same convention as every other admin route
    in this module (structural: Depends(require_admin) in the signature;
    this fake router's `_call` bypasses real FastAPI DI, matching how
    every other route in this file is tested)."""
    src = open(at.__file__, encoding="utf-8").read()
    idx = src.index('@api.post("/admin/assessments/submissions/{submission_id}/correction")')
    snippet = src[idx:idx + 400]
    assert "Depends(require_admin)" in snippet


def test_correction_rejected_on_a_never_awarded_submission(monkeypatch):
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, _ = _build(wallet=wallet)
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        return {"ok": True, "answers": [], "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)
    file = _UploadFile(b"bytes", "image/png")
    submitted = _call(router, "POST", "/student/assessments/submit",
                       assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub_id = submitted["submission"]["submissionId"]

    with pytest.raises(Exception) as exc_info:
        _correct(router, sub_id, corrections=[{"qid": "q1", "correct": True, "points": 0.5}])
    assert "409" in str(exc_info.value) or "already-awarded" in str(exc_info.value)
    assert wallet.calls == 0


def test_positive_correction_credits_exact_net_difference(monkeypatch):
    """Q30 was scored incorrect (0 pts) originally; teacher accepts student
    evidence and marks it correct for 0.5 pts. Net change: +0.5, not a
    re-award of the full 15.5."""
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, pushes = _build(wallet=wallet)
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        # Everything correct except the last question (left unanswered).
        return {"ok": True, "answers": [{"qid": q["qid"], "answer": q["correctAnswer"], "confidence": 0.95}
                                          for q in questions[:-1]], "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)
    file = _UploadFile(b"bytes", "image/png")
    submitted = _call(router, "POST", "/student/assessments/submit",
                       assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub_id = submitted["submission"]["submissionId"]
    assert submitted["submission"]["score"]["pointsEarned"] == 14.5

    awarded = _call(router, "POST", "/admin/assessments/submissions/{submission_id}/award",
                     submission_id=sub_id, admin=_Admin())
    assert awarded["points"] == 14.5
    assert wallet.balances["stu_alice"] == 14.5

    last_qid = asmt["questions"][-1]["qid"]
    result = _correct(
        router, sub_id,
        corrections=[{"qid": last_qid, "correct": True, "points": 0.5,
                      "note": "Handwriting was legible on the original paper."}],
        reason="student_evidence_accepted",
    )
    assert result["ok"] is True
    assert result["duplicate"] is False
    assert result["score"]["pointsEarned"] == 15.0
    assert result["score"]["correct"] == 30
    assert result["award"]["pointsCredited"] == 15.0
    assert result["correction"]["walletAdjustment"] == 0.5
    assert result["correctionVersion"] == 1

    # Exactly ONE additional wallet call, for the NET 0.5 — never a second
    # full re-award.
    assert wallet.calls == 2  # original award + this correction credit
    assert wallet.debit_calls == 0
    assert wallet.balances["stu_alice"] == 15.0

    stored = db[at.COLL_SUBMISSIONS].docs[sub_id]
    assert stored["status"] == "awarded"  # status never leaves "awarded"
    assert stored["correctionState"] == "applied"
    assert stored["correctionVersion"] == 1
    assert stored["score"]["pointsEarned"] == 15.0
    # Original award immutably preserved, distinct from the live award.
    assert stored["originalAward"]["pointsCredited"] == 14.5
    assert stored["originalAward"]["score"]["pointsEarned"] == 14.5

    # Notification sent (nonzero diff), correct "increased" framing.
    assert len(pushes) == 2  # original award notification + correction one
    corr_push = pushes[-1]
    assert "increased" in corr_push["body"]
    assert "+0.5" in corr_push["body"]


# ── negative-correction wallet recovery: capped at zero, never negative ──
# Product decision (explicit, post-audit): a correction reversal must NEVER
# push a student's wallet below zero. Recover only what's actually there;
# record the unrecovered remainder as an honest walletShortfall in the
# audit trail. The academic correction (score/points) always applies in
# full regardless of how much could be recovered.
def _negative_correction_setup(monkeypatch, starting_balance):
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, pushes = _build(wallet=wallet)
    sub_id, questions = _award_full_submission(db, router, wallet)
    assert wallet.balances["stu_alice"] == 15.0
    wallet.balances["stu_alice"] = starting_balance
    return wallet, db, router, pushes, sub_id, questions


def test_negative_correction_sufficient_balance_recovers_in_full(monkeypatch):
    """Balance comfortably covers the reversal — full recovery, zero
    shortfall, exactly one debit call for the exact amount. Also confirms
    the notification body never overclaims a specific deducted amount
    (it already didn't, pre-audit) — still true under capped debits."""
    wallet, db, router, pushes, sub_id, questions = _negative_correction_setup(monkeypatch, 10.0)
    q1 = questions[0]["qid"]
    result = _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
                       reason="question_key_error")
    assert result["score"]["pointsEarned"] == 14.5
    assert result["award"]["pointsCredited"] == 14.5  # academic correction applied in full
    assert result["correction"]["walletAdjustment"] == -0.5
    assert result["correction"]["walletShortfall"] == 0.0
    assert wallet.debit_calls == 1
    assert wallet.balances["stu_alice"] == 9.5  # 10.0 - 0.5, never negative
    corr_push = pushes[-1]
    assert "updated from" in corr_push["body"]
    assert "increased" not in corr_push["body"]


def test_negative_correction_exact_balance_recovers_in_full_lands_on_zero(monkeypatch):
    """Balance equals exactly the reversal amount — full recovery, wallet
    lands precisely on zero (not below)."""
    wallet, db, router, _pushes, sub_id, questions = _negative_correction_setup(monkeypatch, 0.5)
    q1 = questions[0]["qid"]
    result = _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
                       reason="question_key_error")
    assert result["correction"]["walletAdjustment"] == -0.5
    assert result["correction"]["walletShortfall"] == 0.0
    assert wallet.balances["stu_alice"] == 0.0
    assert result["award"]["pointsCredited"] == 14.5  # unaffected by wallet capping


def test_negative_correction_insufficient_balance_caps_at_available_and_records_shortfall(monkeypatch):
    """The student already spent most of it — the wallet is debited ONLY
    down to zero (never negative), and the unrecovered remainder is
    recorded explicitly as walletShortfall. The academic correction still
    applies in full."""
    wallet, db, router, _pushes, sub_id, questions = _negative_correction_setup(monkeypatch, 0.2)
    q1 = questions[0]["qid"]
    result = _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
                       reason="question_key_error")
    assert result["score"]["pointsEarned"] == 14.5  # academic correction: full effect
    assert result["award"]["pointsCredited"] == 14.5  # never blocked by the wallet limitation
    assert result["correction"]["walletAdjustment"] == -0.2  # only what was actually recovered
    assert result["correction"]["walletShortfall"] == 0.3  # 0.5 requested - 0.2 recovered
    assert wallet.debit_calls == 1
    assert wallet.balances["stu_alice"] == 0.0  # never negative

    stored = db[at.COLL_SUBMISSIONS].docs[sub_id]
    assert stored["score"]["pointsEarned"] == 14.5
    assert stored["award"]["pointsCredited"] == 14.5


def test_negative_correction_zero_balance_recovers_nothing_never_calls_debit(monkeypatch):
    """The student's balance is already zero — nothing to recover. No
    wallet.debit() call is even attempted; the shortfall is the full
    requested amount; the academic correction still applies."""
    wallet, db, router, _pushes, sub_id, questions = _negative_correction_setup(monkeypatch, 0.0)
    q1 = questions[0]["qid"]
    result = _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
                       reason="question_key_error")
    assert result["score"]["pointsEarned"] == 14.5
    assert result["award"]["pointsCredited"] == 14.5
    assert result["correction"]["walletAdjustment"] == 0.0
    assert result["correction"]["walletShortfall"] == 0.5
    assert wallet.debit_calls == 0  # never even attempted
    assert wallet.balances["stu_alice"] == 0.0


def test_negative_correction_debit_race_failure_still_applies_the_academic_correction(monkeypatch):
    """A genuine race (balance drops between the read and the atomic
    guarded debit, e.g. a concurrent unrelated spend) must NEVER block or
    roll back the grading correction — it degrades to a full shortfall,
    logged, never silently dropped."""
    wallet, db, router, _pushes, sub_id, questions = _negative_correction_setup(monkeypatch, 5.0)
    wallet.debit_fail = True  # simulates the atomic debit itself failing
    q1 = questions[0]["qid"]
    result = _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
                       reason="question_key_error")
    assert result["ok"] is True  # NOT a 502 — the correction still succeeded
    assert result["score"]["pointsEarned"] == 14.5
    assert result["award"]["pointsCredited"] == 14.5
    assert result["correction"]["walletAdjustment"] == 0.0
    assert result["correction"]["walletShortfall"] == 0.5  # full amount, honestly recorded
    stored = db[at.COLL_SUBMISSIONS].docs[sub_id]
    assert stored["correctionState"] == "applied"  # academic correction DID apply


def test_negative_correction_balance_lookup_failure_still_applies_the_academic_correction(monkeypatch):
    """If even the balance READ fails, the correction must still not be
    blocked — degrades to zero-recoverable, full shortfall."""
    wallet, db, router, _pushes, sub_id, questions = _negative_correction_setup(monkeypatch, 5.0)
    wallet.get_balance_fail = True
    q1 = questions[0]["qid"]
    result = _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
                       reason="question_key_error")
    assert result["ok"] is True
    assert result["award"]["pointsCredited"] == 14.5
    assert result["correction"]["walletAdjustment"] == 0.0
    assert result["correction"]["walletShortfall"] == 0.5
    assert wallet.debit_calls == 0  # never reached — the read itself failed


def test_negative_correction_capped_debit_is_idempotent_on_retry(monkeypatch):
    """Double-click Apply on a capped (partial-recovery) negative
    correction must never debit twice — same clientToken, same result,
    same balance."""
    wallet, db, router, _pushes, sub_id, questions = _negative_correction_setup(monkeypatch, 0.2)
    q1 = questions[0]["qid"]
    first = _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
                      reason="question_key_error", client_token="cap-retry")
    assert first["duplicate"] is False
    assert wallet.debit_calls == 1
    assert wallet.balances["stu_alice"] == 0.0

    second = _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
                       reason="question_key_error", client_token="cap-retry")
    assert second["duplicate"] is True
    assert wallet.debit_calls == 1  # not called again
    assert wallet.balances["stu_alice"] == 0.0  # unchanged
    assert second["correction"]["walletShortfall"] == 0.3  # same recorded shortfall, not recomputed


def test_negative_correction_concurrent_stale_request_rejected_before_touching_wallet(monkeypatch):
    """The CAS concurrency guard applies identically to capped-debit
    corrections — a stale (already-superseded) request is rejected before
    it ever reads the balance or touches the wallet."""
    wallet, db, router, _pushes, sub_id, questions = _negative_correction_setup(monkeypatch, 0.2)
    q1, q2 = questions[0]["qid"], questions[1]["qid"]

    first = _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
                      reason="question_key_error", client_token="race-a", expected_version=0)
    assert first["correctionVersion"] == 1
    lookups_after_first = wallet.balance_lookups

    with pytest.raises(Exception) as exc_info:
        _correct(router, sub_id, corrections=[{"qid": q2, "correct": False, "points": 0.0}],
                 reason="question_key_error", client_token="race-b-stale", expected_version=0)
    assert "409" in str(exc_info.value) or "corrected since you opened it" in str(exc_info.value)
    assert wallet.balance_lookups == lookups_after_first  # never even read the balance
    assert wallet.debit_calls == 1  # only the first correction's debit


def test_negative_correction_audit_trail_records_original_and_current_state(monkeypatch):
    """The correction record preserves enough to reconstruct exactly what
    was requested (originalPoints/correctedPoints), what actually moved
    (walletAdjustment), and what could not be recovered (walletShortfall)
    — append-only, never overwritten by a later correction."""
    wallet, db, router, _pushes, sub_id, questions = _negative_correction_setup(monkeypatch, 0.2)
    q1 = questions[0]["qid"]
    _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
             reason="question_key_error", reason_note="", client_token="audit-1")

    history = _call(router, "GET", "/admin/assessments/submissions/{submission_id}/corrections",
                     submission_id=sub_id, admin=_Admin())
    assert len(history["corrections"]) == 1
    entry = history["corrections"][0]
    assert entry["originalPoints"] == 15.0
    assert entry["correctedPoints"] == 14.5
    assert entry["walletAdjustment"] == -0.2
    assert entry["walletShortfall"] == 0.3
    assert entry["reason"] == "question_key_error"


def test_zero_net_correction_makes_no_wallet_call_and_no_notification(monkeypatch):
    """Mixed correction: Q1 flips wrong->right (+0.5), Q2 flips right->wrong
    (-0.5). Net change is exactly zero -> must move ZERO points and send
    NO notification, even though two individual answers changed."""
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, pushes = _build(wallet=wallet)

    async def fake_extract(media_bytes, content_type, questions):
        # q1 wrong (opposite of correct answer), everything else correct.
        answers = []
        for q in questions:
            if q["qid"] == "q1":
                answers.append({"qid": "q1", "answer": "SHORT" if q["correctAnswer"] == "LONG" else "LONG",
                                 "confidence": 0.95})
            else:
                answers.append({"qid": q["qid"], "answer": q["correctAnswer"], "confidence": 0.95})
        return {"ok": True, "answers": answers, "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)
    asmt = _seed_published_assessment(db)
    file = _UploadFile(b"bytes", "image/png")
    submitted = _call(router, "POST", "/student/assessments/submit",
                       assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub_id = submitted["submission"]["submissionId"]
    assert submitted["submission"]["score"]["pointsEarned"] == 14.5

    awarded = _call(router, "POST", "/admin/assessments/submissions/{submission_id}/award",
                     submission_id=sub_id, admin=_Admin())
    assert awarded["points"] == 14.5

    result = _correct(
        router, sub_id,
        corrections=[
            {"qid": "q1", "correct": True, "points": 0.5, "note": "Evidence accepted."},
            {"qid": "q2", "correct": False, "points": 0.0, "note": "Actually misread by teacher."},
        ],
        reason="teacher_grading_mistake",
    )
    assert result["ok"] is True
    assert result["score"]["pointsEarned"] == 14.5  # net unchanged
    assert result["correction"]["walletAdjustment"] == 0.0
    assert wallet.calls == 1  # only the original award
    assert wallet.debit_calls == 0
    assert len(pushes) == 1  # only the original award push — NO correction push
    # But the audit trail still records the real per-question swap.
    changes = {c["qid"]: c for c in result["correction"]["questionChanges"]}
    assert changes["q1"]["newCorrect"] is True
    assert changes["q2"]["newCorrect"] is False


def test_correction_is_idempotent_on_retry_with_same_client_token(monkeypatch):
    """Double-click Apply / a retried network request with the SAME
    clientToken must never move points twice or create a second audit
    record — this is the wallet-safety hard requirement."""
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, pushes = _build(wallet=wallet)
    sub_id, questions = _award_full_submission(db, router, wallet)
    q1 = questions[0]["qid"]

    first = _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
                      reason="teacher_grading_mistake", client_token="retry-tok")
    assert first["duplicate"] is False
    assert wallet.debit_calls == 1
    assert wallet.balances["stu_alice"] == 14.5

    # Retry #1: same token, submitted again (simulating a refresh/resubmit).
    second = _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
                       reason="teacher_grading_mistake", client_token="retry-tok")
    assert second["duplicate"] is True
    assert wallet.debit_calls == 1  # NOT called again
    assert wallet.balances["stu_alice"] == 14.5  # unchanged

    # Retry #2: press Save a third time.
    third = _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
                      reason="teacher_grading_mistake", client_token="retry-tok")
    assert third["duplicate"] is True
    assert wallet.debit_calls == 1

    corrections_list = _call(router, "GET", "/admin/assessments/submissions/{submission_id}/corrections",
                              submission_id=sub_id, admin=_Admin())
    assert len(corrections_list["corrections"]) == 1  # exactly one audit record, not three
    assert len(pushes) == 2  # original award + exactly one correction notification


def test_three_step_correction_chain_matches_the_audit_required_pattern(monkeypatch):
    """The exact scenario the safety audit demanded, verified with this
    fixture's own real point weights (0.5/question, not an illustrative
    round number): a three-correction chain where each step must diff
    against the ROLLING current state, never the ORIGINAL award, and the
    FINAL net wallet movement from the true original must equal the sum
    of the three individual deltas — never double-counted, never computed
    against a stale baseline.

      Original:  28/30 correct — 14.0 pts
      Correction 1 (fix q1):    29/30 — 14.5 pts  (+0.5)
      Correction 2 (fix q2):    30/30 — 15.0 pts  (+0.5)
      Correction 3 (unfix q1):  29/30 — 14.5 pts  (-0.5)
      Net wallet movement from the ORIGINAL 14.0: +0.5 (NOT +1.5, which
      is what a buggy "always diff against the original" implementation
      would produce for the same three steps: +0.5 +0.5 +0.5 = +1.5).
    """
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, _ = _build(wallet=wallet)
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        # q1 and q2 wrong, everything else (28 questions) correct.
        answers = []
        for q in questions:
            if q["qid"] in ("q1", "q2"):
                answers.append({"qid": q["qid"], "answer": "SHORT" if q["correctAnswer"] == "LONG" else "LONG",
                                 "confidence": 0.95})
            else:
                answers.append({"qid": q["qid"], "answer": q["correctAnswer"], "confidence": 0.95})
        return {"ok": True, "answers": answers, "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)
    file = _UploadFile(b"bytes", "image/png")
    submitted = _call(router, "POST", "/student/assessments/submit",
                       assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub_id = submitted["submission"]["submissionId"]
    assert submitted["submission"]["score"]["correct"] == 28
    assert submitted["submission"]["score"]["pointsEarned"] == 14.0

    awarded = _call(router, "POST", "/admin/assessments/submissions/{submission_id}/award",
                     submission_id=sub_id, admin=_Admin())
    assert awarded["points"] == 14.0
    original_balance = wallet.balances["stu_alice"]

    c1 = _correct(router, sub_id, corrections=[{"qid": "q1", "correct": True, "points": 0.5}],
                  reason="student_evidence_accepted", client_token="chain-1", expected_version=0)
    assert (c1["score"]["correct"], c1["score"]["pointsEarned"]) == (29, 14.5)
    assert c1["correction"]["walletAdjustment"] == 0.5

    c2 = _correct(router, sub_id, corrections=[{"qid": "q2", "correct": True, "points": 0.5}],
                  reason="student_evidence_accepted", client_token="chain-2", expected_version=1)
    assert (c2["score"]["correct"], c2["score"]["pointsEarned"]) == (30, 15.0)
    assert c2["correction"]["walletAdjustment"] == 0.5

    c3 = _correct(router, sub_id, corrections=[{"qid": "q1", "correct": False, "points": 0.0}],
                  reason="teacher_grading_mistake", client_token="chain-3", expected_version=2)
    assert (c3["score"]["correct"], c3["score"]["pointsEarned"]) == (29, 14.5)
    assert c3["correction"]["walletAdjustment"] == -0.5

    # Final state matches the audit's required end state exactly.
    stored = db[at.COLL_SUBMISSIONS].docs[sub_id]
    assert stored["score"]["correct"] == 29
    assert stored["score"]["pointsEarned"] == 14.5
    assert stored["award"]["pointsCredited"] == 14.5

    # The TRUE net movement from the ORIGINAL award is the sum of the
    # three real deltas (+0.5 +0.5 -0.5 = +0.5) — proven directly against
    # the wallet's own running balance, not just the submission's display
    # fields (which a bug could desync from the real ledger).
    assert wallet.balances["stu_alice"] == original_balance + 0.5

    # Original award (before ANY correction) remains immutable throughout
    # the whole chain — never overwritten by correction #2 or #3.
    assert stored["originalAward"]["pointsCredited"] == 14.0
    assert stored["originalAward"]["score"]["correct"] == 28

    history = _call(router, "GET", "/admin/assessments/submissions/{submission_id}/corrections",
                     submission_id=sub_id, admin=_Admin())
    assert len(history["corrections"]) == 3  # append-only — all three preserved
    assert [c["correctedPoints"] for c in history["corrections"]] == [14.5, 15.0, 14.5]
    assert [c["originalPoints"] for c in history["corrections"]] == [14.0, 14.5, 15.0]  # each vs its OWN prior state


def test_stale_correction_version_is_rejected_not_overwritten(monkeypatch):
    """Edge case 19F: a teacher opens correction mode, another correction
    lands first, then the FIRST teacher tries to apply against a now-stale
    expectedVersion. Must be rejected with a clear signal to refresh —
    never silently overwritten."""
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, _ = _build(wallet=wallet)
    sub_id, questions = _award_full_submission(db, router, wallet)
    q1, q2 = questions[0]["qid"], questions[1]["qid"]

    # Teacher A's correction lands first (version 0 -> 1).
    first = _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
                      reason="teacher_grading_mistake", client_token="teacher-a", expected_version=0)
    assert first["correctionVersion"] == 1

    # Teacher B opened the submission before A's correction landed (still
    # thinks expectedVersion is 0) and now tries to apply.
    with pytest.raises(Exception) as exc_info:
        _correct(router, sub_id, corrections=[{"qid": q2, "correct": False, "points": 0.0}],
                 reason="teacher_grading_mistake", client_token="teacher-b", expected_version=0)
    assert "409" in str(exc_info.value) or "corrected since you opened it" in str(exc_info.value)

    # No point movement happened for teacher B's rejected attempt.
    assert wallet.debit_calls == 1  # only teacher A's correction
    stored = db[at.COLL_SUBMISSIONS].docs[sub_id]
    assert stored["correctionVersion"] == 1

    # Teacher B refreshes (reads the new version) and retries — succeeds.
    retried = _correct(router, sub_id, corrections=[{"qid": q2, "correct": False, "points": 0.0}],
                        reason="teacher_grading_mistake", client_token="teacher-b-v2", expected_version=1)
    assert retried["correctionVersion"] == 2
    assert wallet.debit_calls == 2


def test_multiple_sequential_corrections_each_move_only_their_own_net_diff(monkeypatch):
    """Two separate correction sessions on the same submission — each must
    move exactly its own net diff, and the running total must be correct."""
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, _ = _build(wallet=wallet)
    sub_id, questions = _award_full_submission(db, router, wallet)
    q1, q2 = questions[0]["qid"], questions[1]["qid"]

    r1 = _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
                  reason="teacher_grading_mistake", client_token="c1", expected_version=0)
    assert r1["score"]["pointsEarned"] == 14.5
    assert wallet.balances["stu_alice"] == 14.5

    r2 = _correct(router, sub_id, corrections=[{"qid": q2, "correct": False, "points": 0.0}],
                  reason="teacher_grading_mistake", client_token="c2", expected_version=1)
    assert r2["score"]["pointsEarned"] == 14.0
    assert wallet.balances["stu_alice"] == 14.0
    assert r2["correction"]["walletAdjustment"] == -0.5

    # ORIGINAL award (from before EITHER correction) remains immutable —
    # a second correction must not overwrite the snapshot taken at the
    # first one.
    stored = db[at.COLL_SUBMISSIONS].docs[sub_id]
    assert stored["originalAward"]["pointsCredited"] == 15.0
    assert stored["originalAward"]["score"]["pointsEarned"] == 15.0

    history = _call(router, "GET", "/admin/assessments/submissions/{submission_id}/corrections",
                     submission_id=sub_id, admin=_Admin())
    assert len(history["corrections"]) == 2
    assert history["corrections"][0]["correctedPoints"] == 14.5
    assert history["corrections"][1]["correctedPoints"] == 14.0


def test_correction_never_mutates_original_gemini_extraction_or_evidence(monkeypatch):
    """Requirement: original AI evaluation, evidence, and extraction stay
    untouched — a correction is a NEW layer, never a rewrite."""
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, _ = _build(wallet=wallet)
    sub_id, questions = _award_full_submission(db, router, wallet)
    before = db[at.COLL_SUBMISSIONS].docs[sub_id]
    original_extraction = copy.deepcopy(before["originalExtractedAnswers"])
    media_ref = before["mediaRef"]
    media_key = before["mediaKey"]

    q1 = questions[0]["qid"]
    _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
             reason="teacher_grading_mistake", client_token="ev1")

    after = db[at.COLL_SUBMISSIONS].docs[sub_id]
    assert after["originalExtractedAnswers"] == original_extraction  # untouched
    assert after["mediaRef"] == media_ref  # R2 evidence link untouched
    assert after["mediaKey"] == media_key
    # The base (pre-correction) deterministic score is preserved verbatim
    # inside the correction audit record's originalScore.
    history = _call(router, "GET", "/admin/assessments/submissions/{submission_id}/corrections",
                     submission_id=sub_id, admin=_Admin())
    assert history["corrections"][0]["originalPoints"] == 15.0


def test_correction_rejects_unknown_qid_and_requires_valid_reason(monkeypatch):
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, _ = _build(wallet=wallet)
    sub_id, _questions = _award_full_submission(db, router, wallet)

    with pytest.raises(Exception) as exc_info:
        _correct(router, sub_id, corrections=[{"qid": "not_a_real_qid", "correct": True, "points": 1}],
                 reason="teacher_grading_mistake")
    assert "400" in str(exc_info.value) or "matched a known question" in str(exc_info.value)

    with pytest.raises(Exception) as exc_info2:
        _correct(router, sub_id, corrections=[{"qid": "q1", "correct": True, "points": 0.5}],
                 reason="not_a_valid_reason")
    assert "400" in str(exc_info2.value) or "reason must be" in str(exc_info2.value)


def test_correction_points_clamped_to_question_max_never_exceeds_it(monkeypatch):
    """A teacher accidentally typing a huge/negative corrected-points value
    must be clamped into [0, question.points] — never create free points
    or a negative per-question score."""
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, _ = _build(wallet=wallet)
    sub_id, questions = _award_full_submission(db, router, wallet)
    q1 = questions[0]["qid"]  # worth 0.5 pts

    result = _correct(router, sub_id, corrections=[{"qid": q1, "correct": True, "points": 999}],
                       reason="teacher_grading_mistake")
    detail = next(d for d in result["score"]["details"] if d["qid"] == q1)
    assert detail["pointsEarned"] == 0.5  # clamped to the question's own max, not 999


def test_correction_credit_failure_rolls_back_pending_audit_record(monkeypatch):
    """If the wallet call itself throws, the pending correction reservation
    must be rolled back — never left as a phantom "applied" record with
    no matching wallet movement."""
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, _ = _build(wallet=wallet)
    asmt = _seed_published_assessment(db)

    async def fake_extract(media_bytes, content_type, questions):
        # Last question left unanswered -> a real +0.5 correction is
        # possible (a zero-diff correction would never call the wallet
        # at all, which would make this test unable to exercise failure).
        return {"ok": True, "answers": [{"qid": q["qid"], "answer": q["correctAnswer"], "confidence": 0.95}
                                          for q in asmt["questions"][:-1]], "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)
    file = _UploadFile(b"bytes", "image/png")
    submitted = _call(router, "POST", "/student/assessments/submit",
                       assessment_id=asmt["assessmentId"], file=file, student=_Student())
    sub_id = submitted["submission"]["submissionId"]
    _call(router, "POST", "/admin/assessments/submissions/{submission_id}/award",
          submission_id=sub_id, admin=_Admin())
    last_qid = asmt["questions"][-1]["qid"]

    wallet.fail = True
    with pytest.raises(Exception) as exc_info:
        _correct(router, sub_id, corrections=[{"qid": last_qid, "correct": True, "points": 0.5}],
                 reason="student_evidence_accepted")
    assert "502" in str(exc_info.value) or "credit_failed" in str(exc_info.value)

    assert db[at.COLL_CORRECTIONS].docs == {}  # no orphaned pending record
    stored = db[at.COLL_SUBMISSIONS].docs[sub_id]
    assert stored["correctionState"] == "none"  # never touched
    assert stored["score"]["pointsEarned"] == 14.5  # unchanged
    # The atomic version reservation taken before the wallet call must
    # also be released on failure — otherwise a permanently-inflated
    # correctionVersion would make every future expectedVersion check
    # off-by-one for a correction that never actually happened.
    assert stored["correctionVersion"] == 0


def test_concurrent_correction_cas_rejects_the_loser_before_any_wallet_call(monkeypatch):
    """Direct proof of the actual concurrency guard this route relies on:
    two requests that both read correctionVersion=0 at the same moment can
    NEVER both win the atomic find_one_and_update — MongoDB's real
    per-document atomicity (mirrored exactly by this fake) means only ONE
    compare-and-swap from version 0 -> 1 can ever succeed. This tests the
    actual synchronization primitive directly, which is the rigorous way
    to prove mutual exclusion without needing to fabricate true OS-level
    thread interleaving in a single-threaded test process."""
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, _ = _build(wallet=wallet)
    sub_id, _questions = _award_full_submission(db, router, wallet)

    from pymongo import ReturnDocument
    subs = db[at.COLL_SUBMISSIONS]

    # "Teacher A" and "Teacher B" both read correctionVersion=0 before
    # either has written. A's request reaches the CAS first and wins it.
    won_by_a = run(subs.find_one_and_update(
        {"submissionId": sub_id, "correctionVersion": 0},
        {"$set": {"correctionVersion": 1}},
        return_document=ReturnDocument.AFTER,
    ))
    assert won_by_a is not None
    assert won_by_a["correctionVersion"] == 1

    # B's request now attempts the EXACT SAME compare-and-swap (same
    # starting version, since B's read happened before A's write landed)
    # — it must lose, and lose BEFORE ever reaching the wallet.
    lost_by_b = run(subs.find_one_and_update(
        {"submissionId": sub_id, "correctionVersion": 0},
        {"$set": {"correctionVersion": 1}},
        return_document=ReturnDocument.AFTER,
    ))
    assert lost_by_b is None
    assert db[at.COLL_SUBMISSIONS].docs[sub_id]["correctionVersion"] == 1  # only A's write landed


def test_stale_teacher_b_rejected_end_to_end_never_double_adjusts_the_wallet(monkeypatch):
    """Full end-to-end version of the exact scenario the audit demanded:
    Teacher A opens version 0, applies a correction -> version becomes 1.
    Teacher B, still holding version 0 from before A's correction landed,
    attempts to apply a DIFFERENT correction against expectedVersion=0.
    Teacher B must be rejected — never silently overwrite A's correction,
    never produce a second wallet adjustment."""
    _patch_media_storage(monkeypatch)
    wallet = _Wallet()
    db, router, _ = _build(wallet=wallet)
    sub_id, questions = _award_full_submission(db, router, wallet)
    q1, q2 = questions[0]["qid"], questions[1]["qid"]

    a_result = _correct(router, sub_id, corrections=[{"qid": q1, "correct": False, "points": 0.0}],
                         reason="teacher_grading_mistake", client_token="teacher-a", expected_version=0)
    assert a_result["correctionVersion"] == 1
    assert wallet.debit_calls == 1
    balance_after_a = wallet.balances["stu_alice"]

    with pytest.raises(Exception) as exc_info:
        _correct(router, sub_id, corrections=[{"qid": q2, "correct": False, "points": 0.0}],
                 reason="teacher_grading_mistake", client_token="teacher-b-stale", expected_version=0)
    assert "409" in str(exc_info.value) or "corrected since you opened it" in str(exc_info.value)

    # B never reached the wallet at all — no second debit/credit call.
    assert wallet.debit_calls == 1
    assert wallet.calls == 1  # only the ORIGINAL award's credit
    assert wallet.balances["stu_alice"] == balance_after_a

    stored = db[at.COLL_SUBMISSIONS].docs[sub_id]
    assert stored["correctionVersion"] == 1  # NOT bumped again by B's rejected attempt
    # A's correction (q1 -> incorrect) must survive, completely untouched
    # by B's rejected attempt.
    q1_detail = next(d for d in stored["score"]["details"] if d["qid"] == q1)
    assert q1_detail["correct"] is False
    q2_detail = next(d for d in stored["score"]["details"] if d["qid"] == q2)
    assert q2_detail["correct"] is True  # B's change never applied


def test_admin_submission_list_never_fabricates_a_name_for_an_unknown_student(monkeypatch):
    _patch_media_storage(monkeypatch)
    db, router, pushes = _build(wallet=_Wallet())
    asmt = _seed_published_assessment(db)
    # No matching row in db.students at all — the submission's studentId/
    # cleanId genuinely has no resolvable name.

    async def fake_extract(media_bytes, content_type, questions):
        return {"ok": True, "answers": [], "engine": "mock"}
    monkeypatch.setattr(ai, "extract_submission_answers", fake_extract)

    file = _UploadFile(b"bytes", "image/png")
    _call(router, "POST", "/student/assessments/submit",
          assessment_id=asmt["assessmentId"], file=file, student=_Student())

    listed = _call(router, "GET", "/admin/assessments/{assessment_id}/submissions",
                    assessment_id=asmt["assessmentId"], status=None, admin=_Admin())
    assert not listed["submissions"][0].get("studentName")
    # The raw identity codes must still be present as the honest fallback.
    assert listed["submissions"][0]["cleanId"] == "stu094"


def test_real_gas_sync_helper_rounds_fractional_points_and_never_raises():
    """assessment scoring supports 0.5-point increments (0.5/question);
    GAS's legacy points ledger is integer-based (mirrors Speaking Lab's own
    /points/grant contract) — the sync helper must round for THIS leg only
    and never raise, even with no config at all."""
    ok, err = run(at._sync_award_to_gas(
        "stu094", 7.5, gas_url=None, treasury_id=None, treasury_password=None,
    ))
    assert ok is False
    assert err == "gas_sync_not_configured"
